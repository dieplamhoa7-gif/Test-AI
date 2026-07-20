from __future__ import annotations

import html
import re
import time
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore
from reportlab.lib import colors  # type: ignore
from reportlab.lib.enums import TA_CENTER, TA_LEFT  # type: ignore
from reportlab.lib.pagesizes import A4  # type: ignore
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
from reportlab.lib.units import mm  # type: ignore
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle  # type: ignore
from reportlab.pdfgen import canvas  # type: ignore

WORKSPACE = Path(r"C:\Users\HoaD-CVDT\.openclaw\workspace")
TEMP_DIR = WORKSPACE / "temp" / "notebooklm-share"


def _clean(s: str) -> str:
    s = re.sub(r"\s+", " ", str(s or "")).strip()
    return s


def _docx_blocks(docx_path: str | Path) -> list[dict[str, Any]]:
    doc = Document(str(docx_path))
    blocks: list[dict[str, Any]] = []
    for p in doc.paragraphs:
        text = _clean(p.text)
        if text:
            style = (p.style.name or "") if p.style else ""
            blocks.append({"type": "p", "style": style, "text": text})
    for tbl in doc.tables:
        rows: list[list[str]] = []
        for row in tbl.rows:
            vals = [_clean(cell.text) for cell in row.cells]
            if any(vals):
                # Deduplicate repeated merged cells while preserving order.
                compact: list[str] = []
                for v in vals:
                    if not compact or compact[-1] != v:
                        compact.append(v)
                rows.append(compact[:5])
        if rows:
            blocks.append({"type": "table", "rows": rows[:40]})
    return blocks


def _split_sections(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current = {"title": "Tổng quan báo cáo", "items": []}
    for b in blocks:
        if b.get("type") == "p":
            txt = str(b.get("text") or "")
            style = str(b.get("style") or "").lower()
            is_heading = "heading" in style or bool(re.match(r"^(\d+[A-Z]?\.|[IVX]+\.|Mục\s+\d+|Executive|Tin tức|Kỹ thuật|Fundamental|Rủi ro|Kịch bản)", txt, re.I))
            if is_heading and current["items"]:
                sections.append(current)
                current = {"title": txt[:140], "items": []}
            elif is_heading:
                current["title"] = txt[:140]
            else:
                current["items"].append(b)
        else:
            current["items"].append(b)
    if current["items"]:
        sections.append(current)
    return sections[:12]



def _extract_metrics(text: str) -> dict[str, str]:
    metrics: dict[str, str] = {}
    patterns = {
        "Giá": r"Giá tham chiếu[^0-9]{0,80}([0-9]+(?:[.,][0-9]+)?)",
        "Ngày dữ liệu": r"ngày close[:\s]+([0-9]{4}-[0-9]{2}-[0-9]{2})",
        "Volume": r"volume[:\s]+([0-9,.]+)",
        "Risk": r"risk[^0-9]{0,30}([0-5](?:[.,][0-9])?/5)",
        "RSI": r"RSI[^0-9]{0,30}([0-9]+(?:[.,][0-9]+)?)",
        "RankScore": r"rankScore[^0-9]{0,30}([0-9]+(?:[.,][0-9]+)?)",
    }
    for k, pat in patterns.items():
        m = re.search(pat, text, re.I)
        if m:
            metrics[k] = m.group(1)
    return metrics


def _sentiment_counts(text: str) -> dict[str, int]:
    return {
        "Tích cực": len(re.findall(r"Tích cực", text, re.I)),
        "Trung tính": len(re.findall(r"Trung tính", text, re.I)),
        "Tiêu cực": len(re.findall(r"Tiêu cực", text, re.I)),
    }


def _fallback_canvas(c: canvas.Canvas, doc: Any) -> None:
    w, h = A4
    c.saveState()
    c.setFillColor(colors.HexColor("#071B3A"))
    c.rect(0, 0, w, h, stroke=0, fill=1)
    c.setFillColor(colors.HexColor("#0B2A5B"))
    c.circle(w * .86, h * .88, 110, stroke=0, fill=1)
    c.setFillColor(colors.HexColor("#123B7A"))
    c.circle(w * .08, h * .10, 130, stroke=0, fill=1)
    c.setStrokeColor(colors.HexColor("#1DD6FF"))
    c.setLineWidth(.45)
    for i in range(0, 12):
        y = 24 + i * 56
        c.line(18, y, w-18, y)
    for i in range(0, 8):
        x = 22 + i * 72
        c.line(x, 20, x, h-20)
    c.setFillColor(colors.HexColor("#66E7FF"))
    c.setFont("Helvetica-Bold", 8)
    c.drawRightString(w-22, 14, "LHInvestment Model3 fallback • evidence-only from DOCX")
    c.restoreState()

def create_fallback_report_from_docx(docx_path: str, title: str = "LHInvestment Model3 fallback report") -> dict[str, Any]:
    """Create a local NotebookLM-like substitute when NotebookLM quota/rate-limit blocks slides.

    Output is evidence-only from the generated DOCX: a dense navy HTML report plus a PDF version.
    It is not a NotebookLM artifact, but it preserves the report content and gives the web UI a usable
    replacement PDF/HTML link instead of leaving PDF/Slide empty.
    """
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    docx = Path(docx_path)
    if not docx.exists():
        raise FileNotFoundError(docx)
    safe = re.sub(r"[^A-Za-z0-9_-]+", "-", f"{docx.stem}-fallback")[:90].strip("-") or "model3-fallback"
    stamp = time.strftime("%Y%m%d-%H%M%S")
    html_path = TEMP_DIR / f"{stamp}-{safe}.html"
    pdf_path = TEMP_DIR / f"{stamp}-{safe}.pdf"

    blocks = _docx_blocks(docx)
    sections = _split_sections(blocks)

    css = """
    body{margin:0;background:#06172f;color:#eaf6ff;font-family:Inter,Arial,'Segoe UI',sans-serif;line-height:1.35}
    .page{max-width:1180px;margin:0 auto;padding:28px;background:radial-gradient(circle at 86% 8%,#1b78cf55,transparent 22%),linear-gradient(135deg,#071b3a,#0b2a5b 55%,#123b7a)}
    .hero{position:relative;overflow:hidden;border:1px solid #35d9ff;border-radius:24px;padding:24px;margin-bottom:16px;box-shadow:0 0 34px rgba(53,217,255,.26);background:linear-gradient(120deg,#06172fcc,#0b376ecc)}
    .hero:after{content:'';position:absolute;right:28px;top:18px;width:260px;height:92px;background:linear-gradient(135deg,transparent 20%,#3ff3 21%,transparent 22%),repeating-linear-gradient(90deg,#6ff5,#6ff5 4px,transparent 4px,transparent 16px);opacity:.45;border-radius:14px}
    h1{font-size:34px;margin:0;color:#fff;letter-spacing:-.4px}.sub{color:#9feaff;margin-top:8px;max-width:760px}.kpis{display:grid;grid-template-columns:repeat(6,1fr);gap:10px;margin:12px 0 16px}.kpi{background:#071f47;border:1px solid #38dafb99;border-radius:14px;padding:10px}.kpi small{color:#9feaff}.kpi b{display:block;color:#fff;font-size:18px;margin-top:3px}.grid{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:14px}.card{background:linear-gradient(180deg,rgba(8,30,68,.96),rgba(6,23,50,.96));border:1px solid rgba(100,220,255,.5);border-radius:16px;padding:14px;min-height:145px;box-shadow:inset 0 1px rgba(255,255,255,.08)}.card h2{font-size:16px;margin:0 0 8px;color:#7de7ff}.card h2:before{content:'◈ ';color:#ffd166}.card p{font-size:13px;margin:6px 0}.badge{display:inline-block;background:#12d6a0;color:#062035;border-radius:999px;padding:3px 8px;font-weight:800;font-size:12px;margin-right:6px}.neg{background:#ff6b6b}.neu{background:#ffd166}.miniChart{height:42px;border-radius:10px;background:linear-gradient(135deg,#082a5e,#071b3a);position:relative;margin-top:8px;overflow:hidden}.miniChart svg{width:100%;height:42px}table{width:100%;border-collapse:collapse;margin:8px 0;font-size:12px}td,th{border:1px solid rgba(126,231,255,.32);padding:5px;vertical-align:top}tr:nth-child(odd){background:rgba(255,255,255,.04)}.foot{font-size:11px;color:#b6d8ea;margin-top:16px}
    """
    parts = ["<!doctype html><html><head><meta charset='utf-8'><title>", html.escape(title), "</title><style>", css, "</style></head><body><div class='page'>"]
    full_text = " ".join(str(b.get("text") or "") for b in blocks if b.get("type") == "p")
    metrics = _extract_metrics(full_text)
    sent = _sentiment_counts(full_text)
    kpi_html = "".join(f"<div class='kpi'><small>{html.escape(k)}</small><b>{html.escape(str(v))}</b></div>" for k, v in list(metrics.items())[:6])
    if sent:
        kpi_html += "".join(f"<div class='kpi'><small>{html.escape(k)}</small><b>{v}</b></div>" for k, v in sent.items() if v)
    if not kpi_html:
        kpi_html = "<div class='kpi'><small>Nguồn</small><b>DOCX</b></div><div class='kpi'><small>Mode</small><b>Fallback</b></div>"
    spark = "<div class='miniChart'><svg viewBox='0 0 300 42' preserveAspectRatio='none'><polyline fill='none' stroke='#42f5d7' stroke-width='3' points='0,32 35,25 70,28 105,15 140,20 175,10 210,18 245,8 300,14'/><polyline fill='none' stroke='#ffd166' stroke-width='2' opacity='.75' points='0,24 45,28 90,18 135,22 180,14 225,20 300,10'/></svg></div>"
    parts.append(f"<div class='hero'><span class='badge'>Fallback thay NotebookLM</span><h1>{html.escape(title)}</h1><div class='sub'>Tự tạo từ DOCX Model3 khi NotebookLM hết quota/rate-limit. Evidence-only, không lấy dữ liệu ngoài. Thiết kế dạng institutional research infographic.</div>{spark}</div><div class='kpis'>{kpi_html}</div><div class='grid'>")
    for sec in sections:
        parts.append(f"<section class='card'><h2>{html.escape(str(sec['title']))}</h2>")
        count = 0
        for item in sec["items"]:
            if item.get("type") == "table":
                parts.append("<table>")
                for row in item.get("rows", [])[:10]:
                    parts.append("<tr>" + "".join(f"<td>{html.escape(c)[:900]}</td>" for c in row[:4]) + "</tr>")
                parts.append("</table>")
                count += 1
            else:
                txt = html.escape(str(item.get("text") or ""))
                if txt:
                    parts.append(f"<p>{txt[:900]}</p>")
                    count += 1
            if count >= 7:
                break
        parts.append("<div class='miniChart'><svg viewBox='0 0 300 42' preserveAspectRatio='none'><polyline fill='none' stroke='#66e7ff' stroke-width='2.5' points='0,30 45,22 90,27 135,14 180,19 225,11 300,16'/></svg></div></section>")
    parts.append("</div><div class='foot'>Báo cáo thay thế tự động; dùng khi NotebookLM không tạo được PDF/Slide. Nguồn duy nhất: DOCX Model3 đã kiểm tra.</div></div></body></html>")
    html_path.write_text("".join(parts), encoding="utf-8")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.white, fontSize=18, leading=22)
    h_style = ParagraphStyle("h", parent=styles["Heading2"], textColor=colors.HexColor("#66E7FF"), fontSize=12, leading=15, spaceBefore=8)
    p_style = ParagraphStyle("p", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=8.3, leading=10.2, textColor=colors.HexColor("#EAF6FF"))
    doc = SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    flow: list[Any] = [Paragraph(title, title_style), Paragraph("Fallback thay NotebookLM — tạo từ DOCX Model3, evidence-only, infographic style.", p_style), Spacer(1, 4*mm)]
    if metrics:
        flow.append(Table([[Paragraph(f"<b>{html.escape(k)}</b><br/>{html.escape(v)}", p_style) for k, v in list(metrics.items())[:3]], [Paragraph(f"<b>{html.escape(k)}</b><br/>{v}", p_style) for k, v in list(sent.items())]], colWidths=[55*mm,55*mm,55*mm], style=TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#0B2A5B")),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#66E7FF")),("VALIGN",(0,0),(-1,-1),"TOP")])) )
        flow.append(Spacer(1, 4*mm))
    for idx, sec in enumerate(sections, 1):
        if idx == 7:
            flow.append(PageBreak())
        flow.append(Paragraph(html.escape(str(sec["title"])), h_style))
        for item in sec["items"][:6]:
            if item.get("type") == "table":
                data = [[Paragraph(html.escape(c[:450]), p_style) for c in row[:4]] for row in item.get("rows", [])[:8]]
                if data:
                    t = Table(data, repeatRows=1)
                    t.setStyle(TableStyle([("GRID", (0,0), (-1,-1), .25, colors.HexColor("#99BBD1")), ("VALIGN", (0,0), (-1,-1), "TOP"), ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0B5D7A")), ("TEXTCOLOR", (0,0), (-1,-1), colors.white)]))
                    flow.append(t)
            else:
                txt = str(item.get("text") or "")[:1100]
                if txt:
                    flow.append(Paragraph(html.escape(txt), p_style))
        flow.append(Spacer(1, 2*mm))
    doc.build(flow, onFirstPage=_fallback_canvas, onLaterPages=_fallback_canvas)

    return {"ok": True, "fallback": True, "provider": "local_docx_fallback", "html_path": str(html_path), "slide_pdf": str(pdf_path), "pdf_path": str(pdf_path), "warning": "NotebookLM không khả dụng/quota; đã tạo báo cáo thay thế từ DOCX Model3."}
