from __future__ import annotations

import json
import re
import time
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from model3_lhinvestment_context import load_json_record, load_cache_news, load_strategy_records, _derived_indicators_from_history
from vietnamese_text_guard import repair_vietnamese_text, vietnamese_quality_report

BANNED_PATTERNS = [
    r"Trend\s*Pullback",
    r"TrendLH",
    r"Strategy record\s*2",
    r"Trading Dashboard HTML",
    r"Agent lỗi",
    r"HTTPError",
    r"AI nói chuyện",
    r"báo cáo cần chỉnh chu",
    r"prompt",
    r"Kiro",
    r"Codex",
    r"agent",
    r"Model\s*3",
    r"NotebookLM\s*context",
    r"TradingAgents",
    r"Research Pack",
]
IMPORTANT_INDICATORS = ("RSI", "MACD", "histogram", "ADX", "+DI", "-DI", "Bollinger", "bbPercent", "bbUpper", "bbLower", "Ichimoku", "ROC20", "ret5", "Keltner", "stop", "invalid", "rankScore", "buyScore", "riskScore")

# Vietnamese DOCX font skill for Super_LH:
# python-docx's run.font.name alone does not always set the East Asian font slot.
# Microsoft Word may then substitute a bad font and render Vietnamese accents poorly.
# Always set ascii/hAnsi/eastAsia/cs to a Unicode Vietnamese-safe font.
VI_FONT = "Arial"
VI_FONT_ALT = "Calibri"

def _repair_text_quality(text: Any) -> str:
    s = repair_vietnamese_text(text)
    # Fix single-letter pipes from badly stripped table headers.
    s = re.sub(r"\biu kiện\b", "Điều kiện", s, flags=re.I)
    s = re.sub(r"\bK thut\b", "Kỹ thuật", s, flags=re.I)
    glue_fixes = {
        "Kỹ thuậtỷ đồngắn": "Kỹ thuật ngắn",
        "kỹ thuậtỷ đồngắn": "kỹ thuật ngắn",
        "Tóm tắtỷ đồngắn": "Tóm tắt ngắn",
        "tóm tắtỷ đồngắn": "tóm tắt ngắn",
        "Fundamentalvà": "Fundamental và",
        "đầutư": "đầu tư",
        "Kếthoạch": "Kế hoạch",
        "chínhỗ trợong": "chính trong",
        "ỗ trợong": "trong",
        "catalystỷ đồngành": "catalyst ngành",
        "ngànhỷ đồngán lẻ": "ngành bán lẻ",
    }
    for src, dst in glue_fixes.items():
        s = s.replace(src, dst)
    return s


def _force_run_font(run, size: int | None = None, bold: bool | None = None, color: tuple[int, int, int] | None = None, font_name: str = VI_FONT) -> None:
    if bold is not None:
        run.bold = bold
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), font_name)


def _force_doc_fonts(doc: Document, font_name: str = VI_FONT) -> None:
    for style_name in ("Normal", "Body Text", "Table Grid", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
            style.font.name = font_name
            r_pr = style.element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                r_fonts.set(qn(key), font_name)
        except Exception:
            continue


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _cell(cell, text: Any, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(_repair_text_quality(text))
    _force_run_font(r, size=9, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    r = p.add_run(_repair_text_quality(text))
    heading_color = (23, 54, 93) if level == 1 else (145, 85, 145)
    _force_run_font(r, size=15 if level == 1 else 11, bold=True, color=heading_color)


def _is_investor_low_confidence(text: Any) -> bool:
    """Investor-facing report rule: if evidence is missing, omit the analysis instead of printing uncertainty."""
    s = _clean_inline(str(text or ""))
    if not s:
        return True
    low_patterns = (
        r"\bN/A\b",
        r"chưa có",
        r"chưa đủ",
        r"thiếu dữ liệu",
        r"thiếu chỉ tiêu",
        r"cần bổ sung",
        r"cần refresh",
        r"cần chạy lại",
        r"manual review",
        r"không suy đoán",
        r"không bổ sung tin",
        r"chưa đo được",
        r"chưa được kiểm chứng",
    )
    return any(re.search(p, s, re.I) for p in low_patterns)


def _clean_cell(text: Any, max_len: int = 500) -> str:
    s = _clean_inline(str(text or ""))
    if max_len and len(s) > max_len:
        return s[: max_len - 1].rstrip() + "…"
    return s


def _investor_clean_cell(text: Any) -> str:
    s = _clean_inline(str(text or ""))
    # Remove low-confidence clauses inside otherwise useful rows.
    s = re.sub(r"(?i)(?:^|[;,.]\s*)[^;,.]*\b(?:N/A|thiếu dữ liệu|chưa có|chưa đủ|cần bổ sung|cần refresh|manual review|không suy đoán)[^;,.]*(?=$|[;,.])", "", s)
    s = re.sub(r"\s*;\s*;\s*", "; ", s).strip(" ;,.|-")
    return s


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        item = _investor_clean_cell(item)
        if not item or _is_investor_low_confidence(item):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        r = p.add_run("• ")
        _force_run_font(r, size=9, bold=True)
        rr = p.add_run(item)
        _force_run_font(rr, size=9)


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    clean_rows: list[list[str]] = []
    for row in rows:
        vals = [_investor_clean_cell(v) for v in row[: len(headers)]]
        joined = " ".join(vals)
        data_vals = vals[1:] if len(vals) > 1 else vals
        if not any(v for v in data_vals):
            continue
        if _is_investor_low_confidence(joined):
            continue
        clean_rows.append(vals)
    if not clean_rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        _cell(table.rows[0].cells[i], h, True, (255, 255, 255))
        _shade(table.rows[0].cells[i], "1F4E79")
    for row in clean_rows:
        cells = table.add_row().cells
        for i, v in enumerate(row[: len(headers)]):
            _cell(cells[i], v)
    doc.add_paragraph()


def _strip_after_section8(text: str) -> str:
    m = re.search(r"(?im)^\s*##+\s*8[\).\s-]", text)
    if m:
        return text[: m.start()]
    return text


def _drop_banned_blocks(text: str) -> str:
    text = _strip_after_section8(text)
    blocks = re.split(r"\n\s*\n", text)
    keep: list[str] = []
    skip = False
    for block in blocks:
        b = block.strip()
        if not b:
            continue
        if any(re.search(p, b, flags=re.I) for p in BANNED_PATTERNS):
            # Skip entire Trend Pullback/TrendLH/AI-error blocks.
            skip = bool(re.search(r"Trend\s*Pullback|TrendLH|Strategy record\s*2", b, flags=re.I))
            continue
        if skip and re.match(r"^#{1,4}\s", b):
            skip = False
        if skip:
            continue
        keep.append(b)
    return "\n\n".join(keep)


def _clean_inline(text: str) -> str:
    text = _repair_text_quality(text)
    # Never allow raw WordprocessingML/HTML-like markup to leak into visible DOCX cells.
    # A previous build could pass table XML fragments through summary fields, making
    # Word look blank/garbled even though document.xml contained many characters.
    text = re.sub(r"<w:[^>]+>", " ", text)
    text = re.sub(r"</w:[^>]+>", " ", text)
    text = re.sub(r"<[^>]{1,200}>", " ", text)
    text = re.sub(r"\b(?:w:tblPr|w:tblGrid|w:tr|w:tc|w:p|w:r|w:t|w:tcPr|w:rPr|w:pPr)\b", " ", text)
    text = re.sub(r"https?://\S+", "", text)
    text = text.replace("**", "")
    text = text.replace("###", "").replace("##", "").replace("####", "")
    text = re.sub(r"\s+", " ", text).strip(" -|\t")
    return text


ANSI_RE = re.compile(r"\x1b\][^\x07]*(?:\x07|\x1b\\)|\x1b\[[0-?]*[ -/]*[@-~]|\x1b[@-_]")

def _strip_terminal_noise(text: str) -> str:
    s = ANSI_RE.sub("", str(text or ""))
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", s).replace("\r", "\n")
    noise = (
        "Starting session", "Waiting for response", "Thinking", "Responding", "Enter:send", "Shift+Tab",
        "Ctrl+x", "FINAL_RESULT_START", "FINAL_RESULT_END", "START_RESULT", "END_RESULT",
        "TRADINGAGENTS_NEWS_CONTEXT_ERROR", "WEB_SEARCH_CONTEXT", "QUY TẮC BẮT BUỘC",
    )
    lines = []
    for raw in s.splitlines():
        line = raw.strip()
        if not line:
            continue
        if any(n in line for n in noise):
            continue
        if re.fullmatch(r"[─│┃❙◆╭╮╰╯┌┐└┘⠋⠙⠹⠸⠼⠴⠦⠧\s]+", line):
            continue
        lines.append(line)
    return "\n".join(lines)


def _is_grok_news_quality_bad(text: str) -> bool:
    s = _strip_terminal_noise(text)
    q = vietnamese_quality_report(s)
    if int(q.get("replacement_chars", 0)) > 0 or int(q.get("mojibake_markers", 0)) > 0:
        return True
    if int(q.get("unaccented_markers", 0)) >= 8:
        # Grok news can contain URLs/source names/tickers that trip the generic
        # unaccented heuristic. If the section has clear news structure and
        # citations, keep it and let the extractor parse it.
        if re.search(r"(?is)(Ngày|Tiêu đề|Nguồn|URL|Tóm tắt|Tác động).{0,800}(https?://|CafeF|Dân trí|VnExpress|DSC|BSC|VietnamBiz)", s):
            return False
        return True
    if len(re.findall(r"\b(?:trog|ngồn|tn|sosi|điềukiện|kếhoạch|giảmv|phaloãng|tácđộng)\b", s, re.I)) >= 1:
        return True
    return False


def _extract_symbol(task: str) -> str:
    m = re.search(r"\b[A-Z]{2,5}\b", task.upper())
    return m.group(0) if m else "Cổ phiếu"


def _num(v: Any) -> float | None:
    if v in (None, "", "N/A"):
        return None
    try:
        return float(str(v).replace("%", "").replace(",", ".").strip())
    except Exception:
        return None


def _pct_delta(value: Any, base: Any) -> str:
    v, b = _num(value), _num(base)
    if v is None or b in (None, 0):
        return "N/A"
    return f"{(v / b - 1) * 100:+.1f}%"


def _fmt(v: Any, decimals: int = 2) -> str:
    """Format số cho báo cáo: volume có phân cách nghìn, indicator gọn decimals chữ số."""
    n = _num(v)
    if n is None:
        return "N/A"
    if abs(n) >= 10000:
        return f"{n:,.0f}"
    s = f"{n:,.{decimals}f}".rstrip("0").rstrip(".")
    return s or "0"


def _dedupe_lines(lines: list[str], seen: set[str] | None = None) -> list[str]:
    """Loại câu trùng lặp (chuẩn hóa chữ thường, bỏ khoảng trắng) giữa các mục."""
    seen = seen if seen is not None else set()
    out: list[str] = []
    for line in lines:
        key = re.sub(r"\W+", "", str(line).lower())[:160]
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(line)
    return out


def _synthesized_stance(close_v, ma20, ma50, ma200, rsi_v, macd_v, sig_v, volratio_v, adx_v, pdi_v, mdi_v, zone_state, setup_type, news_pos: int, news_neg: int, news_total: int) -> tuple[str, list[list[str]]]:
    """Quan điểm tổng hợp định lượng từ 4 lớp: xu hướng, động lượng, dòng tiền, tin tức.
    Trả (câu quan điểm, bảng điểm thành phần). Không suy đoán khi thiếu dữ liệu."""
    close_n = _num(close_v)
    score = 0
    max_score = 0
    rows: list[list[str]] = []

    # Mỗi lớp chỉ tính vào mẫu số khi CÓ dữ liệu — thiếu dữ liệu không bị tính như tín hiệu xấu.
    trend_pts, trend_max, trend_ev = 0, 0, []
    for label, ma in (("MA20", ma20), ("MA50", ma50), ("MA200", ma200)):
        ma_n = _num(ma)
        if close_n is not None and ma_n:
            trend_max += 1
            if close_n >= ma_n:
                trend_pts += 1
                trend_ev.append(f"trên {label}")
            else:
                trend_ev.append(f"dưới {label}")
    rows.append(["Xu hướng", f"{trend_pts}/{trend_max}" if trend_max else "N/A", "; ".join(trend_ev) or "thiếu dữ liệu MA"])
    score += trend_pts; max_score += trend_max

    mom_pts, mom_max, mom_ev = 0, 0, []
    rsi_n, macd_n, sig_n = _num(rsi_v), _num(macd_v), _num(sig_v)
    if rsi_n is not None:
        mom_max += 1
        if rsi_n >= 50:
            mom_pts += 1; mom_ev.append(f"RSI {rsi_n:.1f} ≥ 50")
        else:
            mom_ev.append(f"RSI {rsi_n:.1f} < 50")
    else:
        mom_ev.append("RSI thiếu")
    if macd_n is not None and sig_n is not None:
        mom_max += 1
        if macd_n >= sig_n:
            mom_pts += 1; mom_ev.append("MACD trên signal")
        else:
            mom_ev.append("MACD dưới signal")
    else:
        mom_ev.append("MACD thiếu")
    rows.append(["Động lượng", f"{mom_pts}/{mom_max}" if mom_max else "N/A", "; ".join(mom_ev)])
    score += mom_pts; max_score += mom_max

    flow_pts, flow_max, flow_ev = 0, 0, []
    vr_n, adx_n = _num(volratio_v), _num(adx_v)
    if vr_n is not None:
        flow_max += 1
        if vr_n >= 1.2:
            flow_pts += 1; flow_ev.append(f"VolumeRatio {vr_n:.2f} ≥ 1.2")
        else:
            flow_ev.append(f"VolumeRatio {vr_n:.2f} < 1.2")
    else:
        flow_ev.append("VolumeRatio thiếu")
    if adx_n is not None:
        flow_max += 1
        if adx_n >= 25:
            flow_pts += 1; flow_ev.append(f"ADX {adx_n:.1f} ≥ 25 (xu hướng có lực)")
        elif adx_n >= 20:
            flow_ev.append(f"ADX {adx_n:.1f} trung tính")
        else:
            flow_ev.append(f"ADX {adx_n:.1f} < 20 (xu hướng yếu)")
    else:
        flow_ev.append("ADX thiếu — không kết luận sức mạnh xu hướng")
    pdi_n, mdi_n = _num(pdi_v), _num(mdi_v)
    if pdi_n is not None and mdi_n is not None:
        flow_max += 1
        if pdi_n >= mdi_n:
            flow_pts += 1; flow_ev.append(f"+DI {pdi_n:.1f} ≥ -DI {mdi_n:.1f}")
        else:
            flow_ev.append(f"-DI {mdi_n:.1f} > +DI {pdi_n:.1f} (bên bán ưu thế)")
    rows.append(["Dòng tiền/sức mạnh", f"{flow_pts}/{flow_max}" if flow_max else "N/A", "; ".join(flow_ev)])
    score += flow_pts; max_score += flow_max

    news_pts, news_max = 0, 0
    if news_total > 0:
        news_max = 2
        if news_pos > news_neg:
            news_pts = 2 if news_neg == 0 else 1
        elif news_pos == news_neg:
            news_pts = 1
        news_ev = f"{news_pos} tích cực / {news_neg} tiêu cực / {news_total - news_pos - news_neg} trung tính trên {news_total} tin"
    else:
        news_ev = "chưa có tin đã kiểm chứng"
    rows.append(["Tin tức/catalyst", f"{news_pts}/{news_max}" if news_max else "N/A", news_ev])
    score += news_pts; max_score += news_max

    ratio = score / max_score if max_score else 0
    if max_score < 4:
        verdict = "CHƯA ĐỦ DỮ LIỆU — nhiều lớp đánh giá thiếu số liệu, cần refresh cache trước khi kết luận"
    elif ratio >= 0.7:
        verdict = "TÍCH CỰC — nền giá và tin tức ủng hộ, ưu tiên theo dõi mua khi có xác nhận volume"
    elif ratio >= 0.5:
        verdict = "TÍCH CỰC THẬN TRỌNG — tín hiệu nghiêng tích cực nhưng chưa đồng thuận, chờ xác nhận"
    elif ratio >= 0.35:
        verdict = "TRUNG TÍNH — tín hiệu trái chiều, chỉ quan sát, không mở vị thế mới"
    else:
        verdict = "THẬN TRỌNG — tín hiệu yếu, ưu tiên quản trị rủi ro/đứng ngoài"
    # Đối chiếu hệ thống V3: cảnh báo quá mua/tránh mua thì không cho verdict tích cực.
    zone_l = str(zone_state or "").lower()
    setup_l = str(setup_type or "").lower()
    overbought = ("quá mua" in zone_l) or ("tránh mua" in setup_l)
    downgraded = False
    if overbought and max_score >= 4 and ratio >= 0.5:
        verdict = "TRUNG TÍNH (hạ bậc do hệ thống V3 cảnh báo vùng quá mua) — chỉ quan sát, không mua đuổi"
        downgraded = True
    recon_note = f"Zone: {zone_state or 'N/A'}; Setup: {setup_type or 'N/A'}"
    recon_note += "; verdict bị hạ 1 bậc do cảnh báo quá mua/tránh mua" if downgraded else ("; cảnh báo quá mua (điểm vốn đã thận trọng)" if overbought else "; không mâu thuẫn lớn với điểm tổng hợp")
    stance = f"{verdict}. Điểm tổng hợp {score}/{max_score} (xu hướng {trend_pts}/{trend_max or '–'}, động lượng {mom_pts}/{mom_max or '–'}, dòng tiền {flow_pts}/{flow_max or '–'}, tin tức {news_pts}/{news_max or '–'})."
    rows.append(["Đối chiếu hệ thống V3", "—", recon_note])
    rows.append(["TỔNG HỢP", f"{score}/{max_score}", verdict])
    return stance, rows


def _price_position_note(close_v: Any, support_v: Any, resistance_v: Any, ma20_v: Any, ma50_v: Any, ma200_v: Any, rsi_v: Any, macd_v: Any, sig_v: Any, volratio_v: Any) -> str:
    close_n = _num(close_v)
    support_n, resistance_n = _num(support_v), _num(resistance_v)
    rsi_n, macd_n, sig_n, vr_n = _num(rsi_v), _num(macd_v), _num(sig_v), _num(volratio_v)
    parts: list[str] = []
    if close_n is not None:
        parts.append(f"Giá tham chiếu {_fmt(close_v)}")
        # Quy ước thống nhất toàn báo cáo: khoảng cách level so với giá hiện tại
        # (hỗ trợ dưới giá mang dấu âm, kháng cự trên giá mang dấu dương).
        if support_n:
            parts.append(f"hỗ trợ {_fmt(support_v)} cách giá {_pct_delta(support_v, close_v)}")
        if resistance_n:
            parts.append(f"kháng cự {_fmt(resistance_v)} cách giá {_pct_delta(resistance_v, close_v)}")
        for label, ma in (("MA20", ma20_v), ("MA50", ma50_v), ("MA200", ma200_v)):
            if _num(ma):
                parts.append(f"giá so với {label} {_fmt(ma)}: {_pct_delta(close_v, ma)}")
    if rsi_n is not None:
        parts.append("RSI trên 50 hỗ trợ động lượng" if rsi_n >= 50 else "RSI dưới 50 cho thấy động lượng chưa xác nhận")
    if macd_n is not None and sig_n is not None:
        parts.append("MACD trên signal là điểm cộng" if macd_n >= sig_n else "MACD dưới signal là điểm trừ")
    if vr_n is not None:
        parts.append("volume xác nhận tốt" if vr_n >= 1.2 else "volume chưa đủ xác nhận breakout")
    return "; ".join(parts) + "." if parts else "Thiếu dữ liệu giá/indicator để định vị chính xác; cần refresh cache trước khi kết luận."


def _load_macro_dashboard_data() -> dict[str, Any]:
    from model3_lhinvestment_context import WORKSPACE as _WS
    candidates = [
        Path(r"C:\Users\HoaD-CVDT\.openclaw\workspace\Vĩ Mô\web\macro_dashboard_data.json"),
        _WS / "Vĩ Mô" / "web" / "macro_dashboard_data.json",
        _WS / "Vi mo" / "web" / "macro_dashboard_data.json",
        Path("Vĩ Mô") / "web" / "macro_dashboard_data.json",
    ]
    for p in candidates:
        try:
            if p.exists():
                data = json.loads(p.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    return data
        except Exception:
            continue
    return {}


def _latest_indicator_value(data: dict[str, Any], keywords: tuple[str, ...], group_keywords: tuple[str, ...] = ()) -> tuple[str, str, str, str, str]:
    best = None
    best_score = -1
    for item in data.get("indicators") or []:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "")
        group = str(item.get("group") or "")
        unit = str(item.get("unit") or "")
        hay = (name + " " + group + " " + str(item.get("category") or "")).lower()
        score = sum(3 for k in keywords if k.lower() in hay) + sum(1 for k in group_keywords if k.lower() in hay)
        # Ưu tiên chỉ số tổng của VN; phạt series ngoại (EU/các nước, đơn vị USD) để không bắt nhầm
        # kiểu "Tăng trưởng GDP thực — EU (Triệu USD)" thay cho GDP Việt Nam.
        if "tổng số" in hay:
            score += 5
        if "usd" in unit.lower():
            score -= 4
        if "cơ cấu" in hay:
            score -= 3
        pts = item.get("points") if isinstance(item.get("points"), list) else []
        pts_2026 = [p for p in pts if isinstance(p, dict) and "2026" in str(p.get("x") or "")]
        if not pts_2026 or score <= 0:
            continue
        if score > best_score:
            best_score = score
            best = (item, pts_2026)
    if not best:
        return ("N/A", "N/A", "N/A", "N/A", "N/A")
    item, pts = best
    first, last = pts[0], pts[-1]
    change = "N/A"
    if _num(first.get("y")) is not None and _num(last.get("y")) is not None:
        change = f"{_num(last.get('y')) - _num(first.get('y')):+.3f}"
    _unit = str(item.get("unit") or "").strip()
    _val_str = f"{last.get('y')}{(' ' + _unit) if _unit else ''} ({last.get('x')})"
    return (str(item.get("name") or ""), _val_str, change, str(item.get("source") or ""), f"{first.get('x')} → {last.get('x')}")


def _extract_const_data_from_html(path: Path) -> dict[str, Any]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return {}
    m = re.search(r"const\s+DATA\s*=\s*(\{.*?\})\s*;", text, re.S)
    if not m:
        return {}
    try:
        data = json.loads(m.group(1))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _load_web_vi_mo_data() -> dict[str, Any]:
    from model3_lhinvestment_context import WORKSPACE as _WS
    candidates = [
        _WS / "stock-news-backend" / "firebase_public" / "web-vi-mo.html",
        Path(r"C:\Users\HoaD-CVDT\.openclaw\workspace\stock-news-backend\firebase_public\web-vi-mo.html"),
    ]
    inbound = Path(r"C:\Users\HoaD-CVDT\.openclaw\media\inbound")
    try:
        candidates = sorted(inbound.glob("Web_Vi_Mo---*.html"), key=lambda x: x.stat().st_mtime, reverse=True) + candidates
    except Exception:
        pass
    for path in candidates:
        data = _extract_const_data_from_html(path)
        if data:
            data["_source_path"] = str(path)
            return data
    return {}


def _macro_label_map(data: dict[str, Any]) -> dict[str, str]:
    indicators = data.get("indicators") if isinstance(data.get("indicators"), dict) else {}
    out = {}
    for k, v in indicators.items():
        out[str(k)] = str((v or {}).get("label") or (v or {}).get("name") or k) if isinstance(v, dict) else str(k)
    return out


def _fmt_corr(v: Any) -> str:
    n = _num(v)
    if n is None:
        return "N/A"
    return f"r={n:+.2f}".replace(".", ",")


def _latest_series_delta(data: dict[str, Any], key: str) -> tuple[Any, Any, Any, float | None, str]:
    months = data.get("months") if isinstance(data.get("months"), list) else []
    ind = (data.get("indicators") or {}).get(key) if isinstance(data.get("indicators"), dict) else None
    series = ind.get("series") if isinstance(ind, dict) and isinstance(ind.get("series"), list) else []
    vals = []
    for i, v in enumerate(series):
        n = _num(v)
        if n is not None:
            period = months[i] if i < len(months) else str(i)
            vals.append((period, n))
    if len(vals) < 2:
        return (None, None, None, None, "N/A")
    prev_p, prev_v = vals[-2]
    last_p, last_v = vals[-1]
    # YoY/rate series are already percentage values, so use percentage-point delta.
    # Absolute-value series (USD/VND, retail, exports, FDI, public investment amount/index) use pct change
    # to avoid turning large VND/USD levels into absurd stock-impact percentages.
    pp_keys = {"credit_yoy", "m2_yoy", "dep_tckt_yoy", "cpi_yoy", "core_yoy", "policy", "deposit", "gdp_growth_ff", "iip_yoy"}
    if key in pp_keys:
        delta = last_v - prev_v
        unit = "điểm %"
    elif prev_v:
        delta = (last_v / prev_v - 1.0) * 100.0
        unit = "% thay đổi"
    else:
        return (last_p, last_v, prev_v, None, "N/A")
    return (last_p, last_v, prev_v, delta, unit)


def _macro_stock_impact_rows_from_web_vi_mo(symbol: str) -> list[list[str]]:
    data = _load_web_vi_mo_data()
    stocks = data.get("stocks") if isinstance(data.get("stocks"), list) else []
    if not stocks:
        return []
    sym = symbol.upper().strip()
    labels = _macro_label_map(data)
    corr_keys = data.get("corrKeys") if isinstance(data.get("corrKeys"), list) else []
    allowed = {"credit_yoy","m2_yoy","dep_tckt_yoy","cpi_yoy","core_yoy","policy","deposit","usdvnd","pubinv_12m","gdp_growth_ff","iip_yoy","retail","exports","fdi_disb"}
    keys = [str(k) for k in corr_keys if str(k) in allowed] or list(allowed)
    target = next((st for st in stocks if isinstance(st, dict) and str(st.get("code") or st.get("symbol") or "").upper() == sym), None)
    rows = []
    if target and isinstance(target.get("macro"), dict):
        pairs=[]
        for k,v in target["macro"].items():
            if str(k) not in allowed: continue
            r=_num(v)
            last_p,last_v,prev_v,delta,unit = _latest_series_delta(data, str(k))
            if r is not None and delta is not None:
                impact = r * delta
                pairs.append((abs(r), str(k), r, delta, impact, last_p, last_v, prev_v, unit))
        for _,k,r,delta,impact,last_p,last_v,prev_v,unit in sorted(pairs, reverse=True)[:3]:
            one_pct_impact = r * 1.0
            latest_direction = "tích cực" if impact > 0 else "tiêu cực"
            sens_direction = "tăng" if one_pct_impact > 0 else "giảm"
            rows.append([
                labels.get(k,k),
                f"Biến động mới nhất {last_p}: {prev_v:.2f} -> {last_v:.2f}; Δ {delta:+.2f} {unit}; tác động hiện tại {impact:+.2f}% ({latest_direction})",
                f"Tương quan {r:+.2f}; nếu {labels.get(k,k)} tăng 1%/1 điểm % thì {sym} ước {sens_direction} {abs(one_pct_impact):.2f}%",
                str(target.get("sec") or "N/A"),
            ])
        if rows: return rows
    for k in ["cpi_yoy","deposit","usdvnd","pubinv_12m","credit_yoy","m2_yoy","exports","fdi_disb"]:
        if k not in keys: continue
        last_p,last_v,prev_v,delta,unit = _latest_series_delta(data, k)
        if delta is None: continue
        vals=[]
        for st in stocks:
            if not isinstance(st, dict) or not isinstance(st.get("macro"), dict): continue
            r=_num(st["macro"].get(k)); code=str(st.get("code") or st.get("symbol") or "").upper()
            if r is not None and code: vals.append((abs(r*delta),r*delta,r,code,str(st.get("sec") or "")))
        vals=sorted(vals, reverse=True)[:5]
        if vals:
            rows.append([labels.get(k,k), f"{last_p}: Δ {delta:+.2f} {unit}; top tác động: " + ", ".join(f"{code} {impact:+.2f}%" for _,impact,_,code,_ in vals), "Xếp theo abs(r × Δ chỉ tiêu) trong DATA.stocks[*].macro.", ", ".join(dict.fromkeys(sec for *_,sec in vals if sec)) or "N/A"])
        if len(rows) >= 8: break
    return rows


def _macro_stock_impact_rows_from_legacy_dashboard(data: dict[str, Any], symbol: str) -> list[list[str]]:
    impacts=[]; sym=symbol.upper()
    for r in ((data.get("vn100Extension") or {}).get("sectorImpact2026") or []):
        if isinstance(r, dict) and sym in [str(t).upper() for t in (r.get("tickers") or [])]: impacts.append(r)
    if not impacts: impacts=[r for r in ((data.get("vn100Extension") or {}).get("sectorImpact2026") or []) if isinstance(r, dict)][:6]
    rows=[]
    for r in impacts[:8]:
        change=r.get("change"); delta="N/A" if change is None else f"{change:+.3f}" if isinstance(change,(int,float)) else str(change)
        tickers=[str(t).upper() for t in (r.get("tickers") or [])]
        rows.append([str(r.get("factor") or r.get("indicator") or "Vĩ mô"), f"{r.get('direction') or ''} {delta}; {r.get('timeframe') or ''}; {r.get('first_value','N/A')} -> {r.get('latest_value','N/A')}", str(r.get("mechanism") or "N/A"), ", ".join((r.get("sectors") or [])[:4]) or "N/A"])
    return rows


def _stock_backend_data_dir() -> Path:
    here = Path(__file__).resolve()
    for parent in [here.parent, *here.parents]:
        candidate = parent / "stock-news-backend"
        if candidate.exists():
            return candidate
    return Path(r"C:\Users\HoaD-CVDT\.openclaw\workspace\stock-news-backend")


def _load_stock_backend_json(*parts: str) -> dict[str, Any] | list[Any] | None:
    root = _stock_backend_data_dir()
    for prefix in ("firebase_public", ""):
        p = root.joinpath(prefix, *parts) if prefix else root.joinpath(*parts)
        try:
            if p.exists():
                return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
    return None


def _records_from_backend(data: Any) -> list[dict[str, Any]]:
    if isinstance(data, dict):
        for key in ("stocks", "items", "data"):
            if isinstance(data.get(key), list):
                return [x for x in data[key] if isinstance(x, dict)]
        if all(isinstance(v, dict) for v in data.values()):
            return [v for v in data.values() if isinstance(v, dict)]
    if isinstance(data, list):
        return [x for x in data if isinstance(x, dict)]
    return []


def _find_symbol_record(records: list[dict[str, Any]], symbol: str) -> dict[str, Any] | None:
    sym = symbol.upper()
    for r in records:
        if str(r.get("symbol") or r.get("code") or "").upper() == sym:
            return r
        syms = r.get("symbols")
        if isinstance(syms, list) and sym in [str(x).upper() for x in syms]:
            return r
    return None


def _fmt_ratio(v: Any) -> str:
    n = _num(v)
    if n is None:
        return "N/A"
    return f"{n:.2f}"


def _ratio_sector_rows(symbol: str) -> list[list[Any]]:
    data = _load_stock_backend_json("data", "valuation_ratios_cache.json")
    records = _records_from_backend(data)
    rec = _find_symbol_record(records, symbol)
    if not rec:
        return []
    sector = str(rec.get("sector") or rec.get("industryCodeLv2") or "N/A")
    peers = [r for r in records if str(r.get("sector") or r.get("industryCodeLv2") or "") == sector]
    rows: list[list[Any]] = []
    for key, label in (("pe", "P/E"), ("pb", "P/B"), ("ps", "P/S")):
        val = _num(rec.get(key))
        vals = [_num(r.get(key)) for r in peers]
        vals = [v for v in vals if v is not None and v > 0]
        if val is None:
            rows.append([label, "N/A", sector, "N/A", "Không có dữ liệu ratio cho mã này trong valuation_ratios_cache."])
            continue
        avg = sum(vals) / len(vals) if vals else None
        med = sorted(vals)[len(vals)//2] if vals else None
        if avg and val > avg * 1.15:
            note = "Cao hơn trung bình ngành — định giá đắt tương đối, cần tăng trưởng/ROE bù đắp."
        elif avg and val < avg * 0.85:
            note = "Thấp hơn trung bình ngành — có thể rẻ tương đối hoặc phản ánh rủi ro/triển vọng yếu."
        elif avg:
            note = "Gần trung bình ngành — cần so thêm tăng trưởng, ROE/biên lợi nhuận và chất lượng tài sản."
        else:
            note = "Thiếu mẫu ngành để so sánh."
        rows.append([label, _fmt_ratio(val), sector, _fmt_ratio(avg), f"Median {_fmt_ratio(med)}. {note}"])
    return rows


def _clean_backend_text(text: Any) -> str:
    s = str(text or "")
    s = s.replace("\\u002F", "/").replace("\\/", "/")
    return _repair_text_quality(s)


def _backend_report_lines(symbol: str, limit: int = 5) -> list[str]:
    sym = symbol.upper()
    lines: list[str] = []
    for file in ("24hmoney_reports.json", "fundamental_signals.json"):
        data = _load_stock_backend_json("data", file)
        for r in _records_from_backend(data):
            syms = [str(r.get("symbol") or "").upper()]
            if isinstance(r.get("symbols"), list):
                syms.extend(str(x).upper() for x in r.get("symbols") or [])
            if sym not in syms:
                continue
            date = str(r.get("report_date") or r.get("date") or "N/A")
            broker = _clean_backend_text(r.get("broker") or r.get("source") or r.get("provider") or "N/A")
            title = _clean_backend_text(r.get("title") or "Báo cáo phân tích")
            summary = _clean_backend_text(r.get("summary") or "").strip()
            target = r.get("target_price")
            extra = f"; target {target:,.0f}" if isinstance(target, (int, float)) else ""
            lines.append(f"{date} - {broker}: {title}{extra}. {summary[:500]}")
            if len(lines) >= limit:
                return lines
    return lines


def _indicator_latest_text(data: dict[str, Any], key: str) -> str:
    months = data.get("months") if isinstance(data.get("months"), list) else []
    ind = (data.get("indicators") or {}).get(key) if isinstance(data.get("indicators"), dict) else None
    if not isinstance(ind, dict): return "N/A"
    series = ind.get("series") if isinstance(ind.get("series"), list) else []
    unit = str(ind.get("unit") or "").strip(); label = str(ind.get("label") or key)
    vals=[]
    for i,v in enumerate(series):
        n=_num(v)
        if n is not None: vals.append((months[i] if i < len(months) else str(i), n))
    if not vals: return f"{label}: chưa có dữ liệu mới nhất"
    period,val=vals[-1]
    return f"{label}: {val:.2f}{unit} ({period})"


def _macro_rows_4a_from_web_vi_mo() -> list[list[str]]:
    data=_load_web_vi_mo_data()
    if not data: return []
    created=str(data.get("createdAt") or data.get("window") or "N/A")
    def join(keys):
        xs=[_indicator_latest_text(data,k) for k in keys]
        return "; ".join(x for x in xs if x != "N/A") or "chưa có dữ liệu mới nhất"
    def latest(key):
        return _latest_series_delta(data, key)
    _, cpi, cpi_prev, cpi_delta, _ = latest("cpi_yoy")
    _, credit, credit_prev, credit_delta, _ = latest("credit_yoy")
    _, m2, m2_prev, m2_delta, _ = latest("m2_yoy")
    _, deposit, deposit_prev, deposit_delta, _ = latest("deposit")
    _, fx, fx_prev, fx_delta, _ = latest("usdvnd")
    _, pub, pub_prev, pub_delta, _ = latest("pubinv_12m")
    growth_note = "Tăng trưởng vẫn là nền hỗ trợ nếu GDP/IIP/bán lẻ/xuất khẩu giữ đà dương; cần theo dõi độ lan tỏa sang lợi nhuận doanh nghiệp."
    inflation_note = "Áp lực lạm phát tăng" if (cpi_delta is not None and cpi_delta > 0) else "Lạm phát hạ nhiệt hoặc chưa tăng thêm"
    inflation_note += "; CPI cao làm tăng rủi ro định giá và kỳ vọng lãi suất."
    money_note_parts=[]
    if credit_delta is not None: money_note_parts.append("tín dụng cải thiện hỗ trợ cầu/dòng tiền" if credit_delta > 0 else "tín dụng chậm lại làm giảm lực đỡ dòng tiền")
    if m2_delta is not None: money_note_parts.append("M2 tăng hỗ trợ thanh khoản" if m2_delta > 0 else "M2 giảm tốc là tín hiệu thận trọng")
    if deposit_delta is not None: money_note_parts.append("lãi suất huy động tăng gây áp lực chiết khấu" if deposit_delta > 0 else "lãi suất huy động giảm hỗ trợ định giá")
    money_note = "; ".join(money_note_parts) or "Theo dõi tín dụng, M2 và lãi suất để đánh giá thanh khoản thị trường."
    fiscal_note = "Đầu tư công tăng là lực đỡ cho hạ tầng/vật liệu/KCN và tâm lý thị trường" if (pub_delta is not None and pub_delta > 0) else "Đầu tư công chưa cho thấy lực đẩy mới, cần theo dõi giải ngân các tháng tới"
    fx_note = "Tỷ giá tăng tạo áp lực ngoại hối, chi phí nhập khẩu và dòng vốn" if (fx_delta is not None and fx_delta > 0) else "Tỷ giá ổn định/giảm giúp giảm áp lực ngoại hối và chiết khấu"
    return [
        ["Tăng trưởng", join(["gdp_growth_ff","iip_yoy","retail","exports","fdi_disb"]), created, growth_note],
        ["Lạm phát", join(["cpi_yoy","core_yoy"]), created, inflation_note],
        ["Tiền tệ / thanh khoản", join(["credit_yoy","m2_yoy","deposit","policy"]), created, money_note],
        ["Tài khóa / đầu tư công", _indicator_latest_text(data,"pubinv_12m"), created, fiscal_note],
        ["Tỷ giá", _indicator_latest_text(data,"usdvnd"), created, fx_note],
    ]


def _macro_tables_for_symbol(symbol: str) -> tuple[list[list[str]], list[list[str]]]:
    data = _load_macro_dashboard_data()
    macro = data.get("macro") if isinstance(data.get("macro"), dict) else {}
    cards = data.get("cards") if isinstance(data.get("cards"), list) else []
    card_by_name = {str(c.get("name") or "").lower(): c for c in cards if isinstance(c, dict)}

    gdp = _latest_indicator_value(data, ("gdp thực", "gdp", "tổng sản phẩm"), ("tài khoản quốc gia",))
    cpi = _latest_indicator_value(data, ("cpi", "lạm phát", "giá tiêu dùng"), ("giá cả",))
    fiscal = _latest_indicator_value(data, ("đầu tư phát triển", "ngân sách", "tài khóa", "nợ công"), ("tài khóa",))
    credit = _latest_indicator_value(data, ("tín dụng", "m2", "cung tiền"), ("tiền tệ",))

    def card_value(name: str, fallback: str = "N/A") -> str:
        c = card_by_name.get(name.lower())
        if not c:
            return fallback
        return f"{c.get('value')} {c.get('unit') or ''}".strip()

    sbv = macro.get("sbv_policy") if isinstance(macro.get("sbv_policy"), dict) else {}
    ib_on = macro.get("ib_on") if isinstance(macro.get("ib_on"), dict) else {}
    fx = macro.get("vcb_fx") if isinstance(macro.get("vcb_fx"), dict) else {}
    macro_rows = _macro_rows_4a_from_web_vi_mo() or [
        ["Vi mô", f"GDP/tăng trưởng: {gdp[1]}; CPI: {cpi[1]}", f"GDP {gdp[4]}; CPI {cpi[4]}", "Giai đoạn chu kỳ: nghiêng về hồi phục/tăng trưởng nếu GDP cải thiện và CPI chưa tạo áp lực lớn; tiếp tục theo dõi PMI/IIP để xác nhận."],
        ["Chính sách tiền tệ", f"Bơm/hút ròng N/A; OMO {sbv.get('omo_rate', card_value('OMO 7 ngày'))}; lãi suất liên ngân hàng qua đêm {(ib_on.get('value') if ib_on.get('value') not in (None, '') else card_value('IB qua đêm'))} %/năm; tái cấp vốn {sbv.get('tai_cap_von', card_value('Tái cấp vốn'))}; tái chiết khấu {sbv.get('tai_chiet_khau', card_value('Tái chiết khấu'))}", str(ib_on.get("date") or sbv.get("effective_date") or data.get("builtAt") or "N/A"), "Nhận xét tăng/giảm: lãi suất liên ngân hàng thấp hỗ trợ thanh khoản; nếu OMO/lãi suất tăng mạnh thì tạo áp lực."],
        ["Tài khóa / đầu tư công", f"Mức giải ngân/đầu tư công: {fiscal[1]}; chính sách tài khóa cụ thể: N/A", fiscal[4], "Ảnh hưởng: tích cực cho hạ tầng/vật liệu/KCN nếu giải ngân tăng; thiếu chính sách cụ thể thì ghi N/A."],
        ["Tín dụng / thanh khoản", f"Tăng trưởng tín dụng/thanh khoản: {credit[1]}; thay đổi {credit[2]}", credit[4], "Tín dụng và thanh khoản cải thiện thì hỗ trợ dòng tiền; thắt lại làm rủi ro chiết khấu tăng."],
        ["Tỉ giá hối đoái", f"DXY N/A; chênh lệch lãi suất N/A; USD/VND {fx.get('usd_transfer') or card_value('USD/VND')}; mua {fx.get('usd_buy', 'N/A')}; bán {fx.get('usd_sell', 'N/A')}", str(fx.get("date") or data.get("builtAt") or "N/A"), "Tỷ giá tăng gây áp lực ngoại hối/dòng vốn; thiếu DXY và spread thì ghi N/A, không suy đoán."],
    ]

    # 4B only: use macro-stock correlation HTML/dashboard guide. Keep 4A above unchanged.
    impact_rows = _macro_stock_impact_rows_from_web_vi_mo(symbol)
    if not impact_rows:
        impact_rows = _macro_stock_impact_rows_from_legacy_dashboard(data, symbol)
    if not impact_rows:
        impact_rows.append(["Vĩ mô", "N/A", "Chưa có dữ liệu tương quan macro trong file HTML/dashboard", "N/A"])
    return macro_rows, impact_rows


def _pick_lines(text: str, keywords: tuple[str, ...], limit: int = 8) -> list[str]:
    out: list[str] = []
    for raw in re.split(r"\n+|(?<=\.)\s+(?=\d+\.|[A-ZĐ])", text):
        line = _clean_inline(raw)
        if len(line) < 12:
            continue
        if any(k.lower() in line.lower() for k in keywords):
            out.append(line[:320])
        if len(out) >= limit:
            break
    return out


def _extract_news_items(text: str, limit: int = 5) -> list[str]:
    """Return investor-facing impactful Grok news items; one `Tin N` section = one row."""
    raw_text = str(text or "").strip()
    clean = re.sub(r"(?is)^.*?(?=##{1,4}\s*Tin\s*1\b|Tin\s*1\s*[-:—])", "", raw_text)
    items: list[str] = []

    def _field(body: str, name_patterns: str) -> str:
        mm = re.search(rf"(?ims)^\s*(?:[-*]\s*)?\*\*(?:{name_patterns})\s*:\*\*\s*(.+?)(?=\n\s*(?:[-*]\s*)?\*\*(?:Ngày|Nguồn|URL|Tóm tắt|Tác động|Tiêu đề)\s*:\*\*|\n\s*---\s*$|\Z)", body)
        return _clean_inline(mm.group(1)) if mm else ""

    def _parse_section(section: str, idx: int) -> str | None:
        sec = section.strip()
        if not sec:
            return None
        head_line = sec.splitlines()[0] if sec.splitlines() else ""
        title = _clean_inline(re.sub(r"(?i)^Tin\s*\d+\s*[-:—]?\s*", "", head_line))
        date = _field(sec, r"Ngày")
        source = _field(sec, r"Nguồn")
        url = _field(sec, r"URL")
        summary = _field(sec, r"Tóm\s*tắt|Tóm\s*t?t")
        impact = _field(sec, r"Tác\s*động|Tác\s*d?ng")
        if not summary:
            # Fallback: first meaningful prose line after heading/source/date.
            prose_lines = []
            for line in sec.splitlines()[1:]:
                if re.match(r"\s*(---|\*\*(Ngày|Nguồn|URL|Tác|Tiêu đề)\s*:)", line, re.I):
                    continue
                val = _clean_inline(line)
                if val:
                    prose_lines.append(val)
            summary = prose_lines[0] if prose_lines else title
        if not impact:
            im = re.search(r"(?is)\*\*(?:Tác\s*động|Tác\s*d?ng)\s*:\*\*\s*(.+)$", sec)
            if im:
                impact = _clean_inline(im.group(1))
        if not (title or summary or impact):
            return None
        return f"Tin {idx} - {title} | Ngày: {date or 'N/A'} | Nguồn/link: {source or 'N/A'} {url or ''} | Tóm tắt: {summary or title} | Tác động: {impact or 'N/A'}"

    # Primary Grok format: `## Tin 1 — ...` separated by horizontal rules.
    parts = re.split(r"(?im)^\s*##{1,4}\s*(?=Tin\s*\d+\b)", clean)
    sections = [p for p in parts if re.match(r"(?is)^\s*Tin\s*\d+\b", p.strip())]
    for sec in sections:
        item = _parse_section(sec, len(items) + 1)
        if item:
            items.append(item[:1400])
        if len(items) >= limit:
            return items

    # Alternate Grok format: `### 1. Title` with fields.
    sec_pat = r"(?ms)^###\s*\d+\.\s*(?P<head>.*?)\n(?P<body>.*?)(?=^---\s*$|^###\s*\d+\.|^##\s+|\Z)"
    for m in re.finditer(sec_pat, raw_text):
        sec = f"Tin {len(items)+1} - {m.group('head')}\n{m.group('body')}"
        item = _parse_section(sec, len(items) + 1)
        if item:
            items.append(item[:1400])
        if len(items) >= limit:
            return items

    # Older flat format: `Tin 1 - ...` without markdown headings.
    pattern = r"(?is)(Tin\s*\d+\s*(?:[-:—][^\n]*)?.*?)(?=(?:\n\s*Tin\s*\d+\s*(?:[-:—]|$))|\Z)"
    for m in re.finditer(pattern, clean):
        item = _parse_section(m.group(1), len(items) + 1)
        if item:
            items.append(item[:1400])
        if len(items) >= limit:
            return items

    # Controlled public-web fallback format produced when Grok provider is down:
    # `1. [2026-..] Title — Source\nURL: ...\nSnippet: ...`
    # Older formatter versions only accepted `Tin 1`, which made section 2 look
    # empty even though RAW WEB RESULTS had usable public/news items.
    def _fallback_news_bucket(title: str, snippet: str) -> str:
        hay = f"{title} {snippet}".lower()
        if re.search(r"tăng vốn|vốn điều lệ|phát hành|cổ phiếu thưởng|chia cổ tức|esop", hay):
            return "capital"
        if re.search(r"lợi nhuận|doanh thu|kqkd|quý\s*(?:1|i\b)|lãi trước thuế|lãi ròng|\blãi\b|kế hoạch kinh doanh", hay):
            return "earnings"
        if re.search(r"thị phần|môi giới|hose|hnx", hay):
            return "market_share"
        if re.search(r"đhđcđ|đại hội cổ đông|cổ đông|hđqt|nhân sự", hay):
            return "agm"
        if re.search(r"khuyến nghị|target|định giá|dự phóng|triển vọng", hay):
            return "broker_view"
        return re.sub(r"\W+", "", hay, flags=re.UNICODE)[:60]

    def _fallback_news_impact(title: str, snippet: str) -> tuple[str, str]:
        hay = f"{title} {snippet}".lower()
        if re.search(r"tăng vốn|vốn điều lệ|phát hành", hay):
            return "Tích cực", "Tăng năng lực margin/tự doanh và vị thế vốn; rủi ro pha loãng nếu lợi nhuận không tăng tương ứng."
        if re.search(r"cổ tức|cổ phiếu thưởng", hay):
            return "Tích cực", "Hỗ trợ tâm lý cổ đông và định giá ngắn hạn; cần đối chiếu ngày chốt quyền và tỷ lệ pha loãng."
        if re.search(r"esop", hay):
            return "Trung tính", "Giữ chân nhân sự nhưng có pha loãng; tích cực chỉ khi đi kèm tăng trưởng lợi nhuận."
        if re.search(r"lợi nhuận|doanh thu|kqkd|lãi trước thuế|lãi ròng|\blãi\b|kế hoạch kinh doanh", hay):
            return "Tích cực", "KQKD/kế hoạch lợi nhuận là catalyst trực tiếp; cần so với kỳ vọng thị trường và nền lợi nhuận năm trước."
        if re.search(r"thị phần|môi giới", hay):
            return "Tích cực", "Thị phần môi giới cải thiện hỗ trợ doanh thu phí và margin; theo dõi tính bền vững qua các quý."
        if re.search(r"đhđcđ|đại hội cổ đông", hay):
            return "Trung tính", "Tin sự kiện quản trị/kế hoạch năm; tác động phụ thuộc nghị quyết cuối cùng và tiến độ thực hiện."
        if re.search(r"khuyến nghị|target|định giá|dự phóng|triển vọng", hay):
            return "Trung tính", "Cung cấp tham chiếu định giá/triển vọng; cần kiểm chứng giả định lợi nhuận và thanh khoản thị trường."
        if re.search(r"sụt giảm|giảm|rủi ro|thanh tra|xử phạt", hay):
            return "Tiêu cực", "Có thể tạo áp lực tâm lý/định giá; cần xác minh mức ảnh hưởng trực tiếp tới SSI."
        return "Trung tính", "Tin liên quan trực tiếp đến SSI; cần đối chiếu cùng giá, thanh khoản và catalyst trước khi hành động."

    web_pat = r"(?ims)^\s*\d+\.\s*\[(?P<date>[^\]]*)\]\s*(?P<title>.+?)\s*(?:\s+—\s+|\s+\|\s+)(?P<source>[^\n]*)\n\s*URL:\s*(?P<url>\S+)\s*\n\s*Snippet:\s*(?P<snippet>.*?)(?=^\s*\d+\.\s*\[|\Z)"
    def _mk_item(title: str, source: str, date: str, url: str, snippet: str) -> str:
        label, impact = _fallback_news_impact(title, snippet)
        summary = snippet or title
        return (
            f"Tin {{idx}} - {title} | Ngày: {date or 'N/A'} | "
            f"Nguồn/link: {source or 'web'} {url} | Nhãn: {label} | Tóm tắt: {summary} | "
            f"Tác động: {impact}"
        )

    seen_buckets: set[str] = set()
    deferred: list[tuple[str, str, str, str, str, str]] = []
    seen_titles: set[str] = set()
    for m in re.finditer(web_pat, raw_text):
        title = _clean_inline(m.group('title'))
        source = _clean_inline(m.group('source'))
        date = _clean_inline(m.group('date'))
        url = _clean_inline(m.group('url'))
        snippet = _clean_inline(m.group('snippet'))
        if not (title or snippet):
            continue
        title_key = re.sub(r"\W+", "", title.lower(), flags=re.UNICODE)[:90]
        if title_key in seen_titles:
            continue
        seen_titles.add(title_key)
        bucket = _fallback_news_bucket(title, snippet)
        rec = (bucket, title, source, date, url, snippet)
        # First pass: one representative per theme so section 2 does not show
        # many rows that all say the same `tăng vốn/ĐHĐCĐ/lợi nhuận` story.
        if bucket in seen_buckets:
            deferred.append(rec)
            continue
        seen_buckets.add(bucket)
        item = _mk_item(title, source, date, url, snippet).replace("Tin {idx}", f"Tin {len(items)+1}")
        items.append(item[:1400])
        if len(items) >= limit:
            return items

    # If public search only found a few themes, fill up with additional distinct
    # titles, but cap repeats per theme so `capital/earnings` do not crowd out
    # market-share/ESOP/broker-view news.
    bucket_counts: dict[str, int] = {}
    for bucket, *_ in deferred:
        bucket_counts.setdefault(bucket, 0)
    for bucket in seen_buckets:
        bucket_counts[bucket] = 1
    priority = {"market_share": 0, "broker_view": 1, "agm": 2, "capital": 3, "earnings": 4}
    deferred.sort(key=lambda r: (bucket_counts.get(r[0], 0), priority.get(r[0], 9)))
    for bucket, title, source, date, url, snippet in deferred:
        if bucket_counts.get(bucket, 0) >= 2:
            continue
        item = _mk_item(title, source, date, url, snippet).replace("Tin {idx}", f"Tin {len(items)+1}")
        items.append(item[:1400])
        bucket_counts[bucket] = bucket_counts.get(bucket, 0) + 1
        if len(items) >= limit:
            return items

    return items


def _news_rows_from_items(items: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in items:
        text = _clean_inline(item)
        lm = re.search(r"(?is)(?:^|[;|])\s*Nhãn\s*:\s*(Tích cực|Tiêu cực|Trung tính)", text)
        label = lm.group(1) if lm else ("Tích cực" if re.search(r"tích cực|positive|bull|\btăng\b", text, re.I) else ("Tiêu cực" if re.search(r"tiêu cực|negative|bear|\bgiảm\b", text, re.I) else "Trung tính"))
        # Output table requested by Hòa Đại ka: only 3 columns — Nhãn, Tóm tắt tin, Tác động.
        summary = text
        impact = "N/A"
        m = re.search(r"(?i)(Tác động\s*:\s*.+)$", text)
        if m:
            impact = _clean_inline(m.group(0))
            summary = _clean_inline(text[:m.start()])
        else:
            m = re.search(r"(?i)(cổ phiếu\s+(?:tăng|giảm|trung tính).+)$", text)
            if m:
                impact = _clean_inline(m.group(1))
                summary = _clean_inline(text[:m.start()])
        summary = re.sub(r"^Tin\s*\d+\s*[—:-]\s*", "", summary, flags=re.I)
        # Hòa Đại ka requested: the news-summary column must contain only the
        # content after "Tóm tắt:" and must not include title/date/source nor the
        # literal "Tóm tắt" label.
        sm = re.search(r"(?is)(?:^|[;|])\s*Tóm\s*tắt\s*:\s*(.+)$", summary)
        if sm:
            summary = _clean_inline(sm.group(1))
        summary = re.sub(r"\s*\|\s*", "; ", summary)
        impact = re.sub(r"\s*\|\s*", "; ", impact)
        rows.append([label, summary[:520], impact[:260]])
    return rows or [["N/A", "Chưa có tin đã kiểm chứng trong kỳ này", "Không bổ sung tin từ nguồn chưa xác thực; cần chạy lại bước Grok news."]]


def _looks_like_bad_ai_text(content: str) -> bool:
    s = str(content or "")
    if not s.strip():
        return True
    if "Agent lỗi" in s or "TimeoutError" in s or "ConnectionError" in s:
        return True
    if "<html" in s.lower():
        return True
    # Grok bridge sometimes leaks Google redirect/base64-like blobs; never put that into DOCX.
    if re.search(r"[A-Za-z0-9_-]{80,}\?oc=", s):
        return True
    if re.search(r"(?:[A-Za-z0-9+/]{60,}=*){2,}", s):
        return True
    q = vietnamese_quality_report(s)
    if int(q.get("replacement_chars", 0)) > 0 or int(q.get("mojibake_markers", 0)) >= 2:
        return True
    # Severe no-accent/missing-letter provider output like "PHN 5  Kịch bản U T" should be excluded.
    if int(q.get("unaccented_markers", 0)) >= 12:
        return True
    return False


def _get_feed_text(state: dict[str, Any], labels: tuple[str, ...], *, allow_bad: bool = False) -> str:
    chunks = []
    for post in state.get("feed", []):
        name = str(post.get("name") or post.get("agent") or "")
        if any(label.lower() in name.lower() for label in labels):
            content = str(post.get("content") or "")
            if not allow_bad and _looks_like_bad_ai_text(content):
                continue
            chunks.append(content)
    return _drop_banned_blocks("\n\n".join(chunks))


def write_model3_docx(task: str, state: dict[str, Any], path: str | Path) -> str:
    symbol = _extract_symbol(task)
    # Section 2 ownership is GrokX only. Do not silently replace GrokX news with Kiro/web/cache output.
    news_text = _strip_terminal_noise(_repair_text_quality(_get_feed_text(state, ("GrokX News & Impact",), allow_bad=True)))
    if _is_grok_news_quality_bad(news_text):
        news_text = "Chưa có bản tin Grok đủ sạch để đưa vào báo cáo. Cần chạy lại Grok news hoặc dùng nguồn web đã kiểm chứng."
    analysis_text = _repair_text_quality(_get_feed_text(state, ("Research", "Analysis", "Codex TA", "Indicator", "Fundamental")))
    scenario_text = _repair_text_quality(_get_feed_text(state, ("Deep Investment Scenario", "Kịch bản")))
    bull_bear_text = _repair_text_quality(_get_feed_text(state, ("Bull/Bear/Catalyst", "Bull", "Bear", "Catalyst")))
    risk_text = _repair_text_quality(_get_feed_text(state, ("Risk & Viewpoint", "Rủi ro")))
    followup_text = _repair_text_quality(_get_feed_text(state, ("Follow-up Plan", "Kế hoạch theo dõi")))
    all_text = _drop_banned_blocks("\n\n".join([news_text, analysis_text, scenario_text, bull_bear_text, risk_text, followup_text]))

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    _force_doc_fonts(doc, VI_FONT)
    doc.styles["Normal"].font.name = VI_FONT
    doc.styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"BÁO CÁO PHÂN TÍCH CỔ PHIẾU — {symbol}")
    _force_run_font(r, size=18, bold=True, color=(23, 54, 93))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run("Báo cáo dành cho nhà đầu tư | Tổng hợp tin tức, kỹ thuật, định giá và rủi ro")
    _force_run_font(rr, size=10, color=(89, 89, 89))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(time.strftime("Thời gian tạo: %Y-%m-%d %H:%M:%S"))
    _force_run_font(rr, size=8, color=(128, 128, 128))

    summary = _pick_lines(all_text, ("sentiment", "giá", "EOD", "rủi ro", "WATCH", "AVOID", "Bullish", "Bearish", "Neutral", "hỗ trợ", "kháng cự", "MACD", "RSI"), 9)
    # Use the latest available price/EOD at report generation as the anchor for all downstream judgment.
    _, _price_eod = load_json_record(symbol, [r"data\v3_full_indicator_cache_v2.json", r"data\eod_all_stocks_hose_hnx.json", r"firebase_public\data\eod_all_stocks_hose_hnx.json", r"data\market_data.json"])
    _price_derived = _derived_indicators_from_history(symbol) or {}
    def _price_val(*names):
        for src in (_price_eod if isinstance(_price_eod, dict) else {}, _price_derived):
            for name in names:
                v = src.get(name)
                if v not in (None, ""):
                    return v
        return "N/A"
    report_price = _price_val('close','Close','price')
    report_date = _price_val('date','tradingDate','time','updatedAt')
    report_volume = _price_val('volume','Volume')

    # ---- Nạp dữ liệu chỉ báo TRƯỚC mục 1 để quan điểm tổng hợp có số liệu thật ----
    # Deterministic fallback from LHInvestment cache/OHLCV so report does not lose indicators when LLM output is thin.
    _, eod = load_json_record(symbol, [r"data\v3_full_indicator_cache_v2.json", r"data\eod_all_stocks_hose_hnx.json", r"firebase_public\data\eod_all_stocks_hose_hnx.json", r"data\market_data.json"])
    _, rs = load_json_record(symbol, [r"data\rs_levels_vn100_cache.json", r"data\rs_levels_hsx_all_cache.json", r"firebase_public\data\rs_levels_hsx_all_cache.json"])
    derived = _derived_indicators_from_history(symbol) or {}
    # Flatten cache V3 ({price, date, indicators{...}, rs{...}}) để không bỏ sót dữ liệu phân tích nested.
    eod_flat: dict[str, Any] = {}
    if isinstance(eod, dict):
        eod_flat.update(eod)
        for _sub in ("indicators", "rs"):
            if isinstance(eod.get(_sub), dict):
                for _k, _v in eod[_sub].items():
                    eod_flat.setdefault(_k, _v)
        if eod_flat.get("close") in (None, "") and eod_flat.get("price") not in (None, ""):
            eod_flat["close"] = eod_flat["price"]
    def val(*names):
        for src in (eod_flat, derived):
            for name in names:
                v = src.get(name)
                if v not in (None, ""):
                    return v
        return "N/A"
    close_v = val('close','Close','price')
    ma10, ma20, ma50, ma100, ma200 = val('ma10','MA10'), val('ma20','MA20'), val('ma50','MA50'), val('ma100','MA100'), val('ma200','MA200')
    rsi_v = val('rsi14','RSI14','rsi')
    macd_v, sig_v, hist_v = val('macd','MACD'), val('macdSignal','MACD_signal','signal'), val('macdHist','MACD_hist','histogram')
    adx_v, pdi_v, mdi_v = val('adx14','ADX','ADX14'), val('plusDI','diPlus','plusDi'), val('minusDI','diMinus','minusDi')
    vol_v, avgvol_v, volratio_v = val('volume','Volume'), val('avgVol20','avgVolume20'), val('volumeRatio')
    # Nhất quán VolumeRatio: nếu có đủ volume & avgvol thì tính lại; chỉ giữ cache khi khớp (±0.1).
    # Tránh trộn 2 nguồn khác phiên làm ratio mâu thuẫn với chính volume/avgvol đang hiển thị.
    _vol_n, _avg_n = _num(vol_v), _num(avgvol_v)
    if _vol_n is not None and _avg_n:
        _vr_calc = round(_vol_n / _avg_n, 2)
        if _num(volratio_v) is None or abs(_vr_calc - _num(volratio_v)) > 0.1:
            volratio_v = _vr_calc
    bb_u, bb_m, bb_l, bbp = val('bbUpper','bollingerUpper','BB_upper'), val('bbMid','bbMiddle','BB_mid'), val('bbLower','bollingerLower','BB_lower'), val('bbPercent')
    rs_s = rs.get('activeSupportDay') if isinstance(rs, dict) else val('supportLevelsDay','support')
    rs_r = rs.get('activeResistanceDay') if isinstance(rs, dict) else val('resistanceLevelsDay','resistance')
    if rs_s in (None, "", "N/A"):
        rs_s = val('supportLevelsDay','support')
    if rs_r in (None, "", "N/A"):
        rs_r = val('resistanceLevelsDay','resistance')
    anchor_note = _price_position_note(close_v, rs_s, rs_r, ma20, ma50, ma200, rsi_v, macd_v, sig_v, volratio_v)

    # Tin tức + đếm sentiment để đưa vào quan điểm tổng hợp.
    news_lines = _extract_news_items(news_text, 8)
    news_rows = _news_rows_from_items(news_lines) if news_lines else []
    news_pos = sum(1 for r in news_rows if r and r[0] == "Tích cực")
    news_neg = sum(1 for r in news_rows if r and r[0] == "Tiêu cực")
    stance, stance_rows = _synthesized_stance(close_v, ma20, ma50, ma200, rsi_v, macd_v, sig_v, volratio_v, adx_v, pdi_v, mdi_v, eod_flat.get("zoneState"), eod_flat.get("setupType"), news_pos, news_neg, len(news_rows))

    _heading(doc, "1. Tóm tắt nhanh", 1)
    _table(doc, ["Hạng mục", "Thông tin chính"], [
        ["Mã cổ phiếu", symbol],
        ["Giá tham chiếu khi xuất báo cáo", f"{_fmt(report_price)} — ngày close: {report_date}; volume: {_fmt(report_volume)}" + (f" (phiên {(_price_derived or {}).get('asOfDate')})" if (isinstance(_price_eod, dict) and _price_eod.get('volume') in (None, '') and str((_price_derived or {}).get('asOfDate') or '') not in ('', str(report_date))) else "")],
        ["Quan điểm tổng hợp", stance],
        ["Vùng giá quan trọng", f"Hỗ trợ {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)} so với giá); kháng cự {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)} so với giá)"],
        ["Lưu ý", "Mọi nhận định kỹ thuật/risk/trigger bên dưới lấy giá tham chiếu khi xuất báo cáo làm mốc; số liệu thiếu ghi N/A, không suy đoán."],
    ])
    _heading(doc, "1B. Điểm tổng hợp theo 4 lớp", 2)
    _table(doc, ["Lớp đánh giá", "Điểm", "Bằng chứng"], stance_rows)
    # Bullet tóm tắt: loại câu trùng với tin tức mục 2 và trùng lẫn nhau.
    seen_keys: set[str] = set()
    _dedupe_lines(news_lines, seen_keys)  # đăng ký tin tức trước để summary không lặp lại
    summary_extra = _dedupe_lines([s for s in summary if not re.match(r"(?i)^\s*Tin\s*\d", s)], seen_keys)
    _bullets(doc, summary_extra[:5])

    _heading(doc, "2. Tin tức liên quan — tối thiểu 5 tin tác động tới cổ phiếu", 1)
    if news_rows:
        _table(doc, ["Nhãn", "Tóm tắt tin", "Tác động"], news_rows[:8])
    else:
        # Do not let an empty Grok result disappear from the PDF. Show Grok's
        # conclusion/diagnostic so the user can see whether Grok searched but
        # found no usable news, or whether the runtime failed.
        raw_grok_note = _clean_cell(news_text, 1400)
        if raw_grok_note:
            _bullets(doc, ["Grok 9router API đã chạy nhưng không xuất được danh sách Tin 1/Tin 2 đủ điều kiện:", raw_grok_note])
    if len(news_lines) < 5:
        _bullets(doc, ["Nguồn tin tức chưa đủ 5 tin đã kiểm chứng tại thời điểm xuất báo cáo; chỉ dùng tin đã xác thực ở trên, không bổ sung tin chưa kiểm chứng để tránh sai lệch."])

    _heading(doc, "3. LHInvestment Indicator Matrix — chỉ báo bắt buộc", 1)
    ta_lines = _pick_lines(analysis_text, IMPORTANT_INDICATORS, 22)
    rows = []
    rows.extend([
        ["RSI14", f"RSI14 {val('rsi14','RSI14','rsi')}. Trên/dưới 50 là mốc xác nhận động lượng quan trọng."],
        ["MACD", f"MACD {val('macd','MACD')}; signal {val('macdSignal','MACD_signal','signal')}; histogram {val('macdHist','MACD_hist','histogram')}."],
        ["ADX/DI", f"ADX14 {val('adx14','ADX','ADX14')}; +DI {val('plusDI','diPlus','plusDi')}; -DI {val('minusDI','diMinus','minusDi')}."],
        ["Bollinger", f"Upper {val('bbUpper','bollingerUpper','BB_upper')}; mid {val('bbMid','bbMiddle','BB_mid')}; lower {val('bbLower','bollingerLower','BB_lower')}; %B {val('bbPercent')}."],
        ["ROC20", f"ROC20 {val('roc20','ROC20')}. ROC dương cho thấy đà 20 phiên cải thiện; ROC âm là suy yếu."],
        ["Return 5 phiên", f"ret5 {val('ret5','return5')}. Dùng để đo xung lực rất ngắn hạn."],
        ["Hỗ trợ/kháng cự EOD", f"Support {val('supportLevelsDay','support')}; Resistance {val('resistanceLevelsDay','resistance')}."],
    ])
    if isinstance(rs, dict):
        rows.extend([
            ["RS status", f"Status {rs.get('srStatusDay')}; stop/invalid {rs.get('stopLossDay') or rs.get('invalidDay') or 'N/A'}"],
            ["Donchian/Keltner", f"Donchian high/low {rs.get('donchianHighDay')}/{rs.get('donchianLowDay')}; Keltner {rs.get('keltnerUpperDay')}/{rs.get('keltnerLowerDay')}"],
        ])
    for line in ta_lines:
        key = next((k for k in IMPORTANT_INDICATORS if k.lower() in line.lower()), "Chỉ báo")
        pair = [key, line]
        if pair not in rows:
            rows.append(pair)

    _heading(doc, "3B. Bốn cặp chỉ báo LHInvestment — nhận định riêng", 2)
    # Câu hành động sinh ĐỘNG theo trạng thái hiện tại, không dùng câu tĩnh mâu thuẫn với số liệu.
    _c_n, _ma20_n, _ma50_n = _num(close_v), _num(ma20), _num(ma50)
    if _c_n is not None and _ma20_n and _ma50_n:
        if _c_n >= _ma20_n and _c_n >= _ma50_n:
            _ma_action = f"giá đã trên MA20/MA50 — giữ quan điểm, chỉ nâng tỷ trọng khi vượt kháng cự {_fmt(rs_r)} có volume"
        elif _c_n >= _ma20_n:
            _ma_action = f"giá đã trên MA20 nhưng còn dưới MA50 {_fmt(ma50)} ({_pct_delta(close_v, ma50)}) — cần đóng cửa trên MA50 trước khi nâng tỷ trọng"
        else:
            _ma_action = f"giá dưới MA20 {_fmt(ma20)} — đứng ngoài, chờ lấy lại MA20/MA50"
    else:
        _ma_action = "thiếu dữ liệu MA — chờ refresh cache trước khi hành động"
    _rsi_n, _hist_n = _num(rsi_v), _num(hist_v)
    if _rsi_n is not None and _hist_n is not None:
        if _rsi_n >= 50 and _hist_n > 0:
            _mom_action = f"RSI {_fmt(rsi_v)} và histogram đã dương — giữ quan điểm khi RSI còn trên 50; cảnh giác nếu RSI rơi lại dưới 50"
        elif _rsi_n >= 50:
            _mom_action = f"RSI {_fmt(rsi_v)} trên 50 nhưng histogram chưa dương — chờ histogram mở rộng dương trước khi xem là tín hiệu mua mạnh"
        else:
            _mom_action = "chờ RSI vượt 50 và histogram dương rõ trước khi xem là tín hiệu mua mạnh"
    else:
        _mom_action = "thiếu RSI/MACD — chờ refresh cache"
    pair_rows = [
        [
            "Cặp 1 — Xu hướng/MA",
            f"Close {_fmt(close_v)}; MA10 {_fmt(ma10)} ({_pct_delta(close_v, ma10)}); MA20 {_fmt(ma20)} ({_pct_delta(close_v, ma20)}); MA50 {_fmt(ma50)} ({_pct_delta(close_v, ma50)}); MA100 {_fmt(ma100)} ({_pct_delta(close_v, ma100)}); MA200 {_fmt(ma200)} ({_pct_delta(close_v, ma200)})",
            f"Kết luận theo giá hiện tại: {anchor_note} Hành động: {_ma_action}; không mua đuổi sát kháng cự {_fmt(rs_r)}.",
        ],
        [
            "Cặp 2 — Động lượng",
            f"RSI14 {_fmt(rsi_v)}; MACD {_fmt(macd_v, 4)}; signal {_fmt(sig_v, 4)}; hist {_fmt(hist_v, 4)}; ROC20 {_fmt(val('roc20','ROC20'))}; ret5 {_fmt(val('ret5','return5'))}",
            f"Kết luận: RSI {'ủng hộ' if (_num(rsi_v) is not None and _num(rsi_v) >= 50) else 'chưa ủng hộ'} xu hướng tăng; MACD {'xác nhận' if (_num(macd_v) is not None and _num(sig_v) is not None and _num(macd_v) >= _num(sig_v)) else 'chưa xác nhận'} động lượng. Hành động: {_mom_action}.",
        ],
        [
            "Cặp 3 — Sức mạnh xu hướng/dòng tiền",
            f"ADX {_fmt(adx_v)}; +DI {_fmt(pdi_v)}; -DI {_fmt(mdi_v)}; Volume {_fmt(vol_v)}; AvgVol20 {_fmt(avgvol_v)}; VolumeRatio {_fmt(volratio_v)}",
            f"Kết luận: ADX {('mạnh (xu hướng có lực)' if _num(adx_v) >= 25 else ('trung tính' if _num(adx_v) >= 20 else 'yếu (xu hướng nhiễu)')) if _num(adx_v) is not None else 'thiếu dữ liệu — không kết luận sức mạnh xu hướng'}; "
            f"{('+DI trên -DI, bên mua chiếm ưu thế' if _num(pdi_v) >= _num(mdi_v) else '-DI trên +DI, bên bán chiếm ưu thế') if (_num(pdi_v) is not None and _num(mdi_v) is not None) else 'DI thiếu dữ liệu'}; "
            f"dòng tiền {'đã vượt bình quân' if (_num(volratio_v) is not None and _num(volratio_v) >= 1.2) else 'chưa xác nhận đủ'}. Hành động: breakout chỉ có giá trị khi VolumeRatio > 1.2 và +DI duy trì trên -DI.",
        ],
        [
            "Cặp 4 — Biên dao động & vùng giá",
            f"BB upper/mid/lower {_fmt(bb_u)}/{_fmt(bb_m)}/{_fmt(bb_l)}; %B {_fmt(bbp)}; Support {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)}); Resistance {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)})",
            f"Kết luận: hỗ trợ {_fmt(rs_s)} cách giá {_pct_delta(rs_s, close_v)}, kháng cự {_fmt(rs_r)} cách giá {_pct_delta(rs_r, close_v)}"
            + (f"; %B {_fmt(bbp)} sát band trên — vùng quá mua ngắn hạn, hạn chế mua đuổi" if (_num(bbp) is not None and _num(bbp) >= 0.9) else (f"; %B {_fmt(bbp)} sát band dưới — vùng quá bán, theo dõi nhịp hồi" if (_num(bbp) is not None and _num(bbp) <= 0.1) else ""))
            + ". Hành động: nếu giá dưới/tiệm cận hỗ trợ thì ưu tiên quản trị rủi ro; nếu vượt kháng cự với volume xác nhận thì mới nâng xác suất bull case.",
        ],
    ]
    _table(doc, ["Cặp chỉ báo", "Dữ liệu chính", "Nhận định riêng"], pair_rows)

    # 3C: tín hiệu hệ thống V3/LHInvestment sẵn có trong cache — evidence-only, không suy đoán.
    _heading(doc, "3C. Tín hiệu hệ thống V3/LHInvestment", 2)
    ichi = eod_flat.get("ichimoku") if isinstance(eod_flat.get("ichimoku"), dict) else {}
    div = eod_flat.get("divergence") if isinstance(eod_flat.get("divergence"), dict) else {}
    v3_rows: list[list[str]] = []
    def _v3(label, value, note=""):
        if value not in (None, "", [], {}, "N/A"):
            v3_rows.append([label, str(value), note])
    _v3("Xu hướng hiệu lực", eod_flat.get("effectiveTrend"), str(eod_flat.get("trendReason") or "")[:220])
    _v3("Sức mạnh xu hướng", eod_flat.get("trendStrength"), f"Cấu trúc thị trường ngày: {eod_flat.get('marketStructureDay') or 'N/A'}")
    _v3("Setup hệ thống", eod_flat.get("setupType"), f"Zone: {eod_flat.get('zoneState') or 'N/A'}; trạng thái volume: {eod_flat.get('volumeState') or 'N/A'}")
    def _score_str(x):
        if isinstance(x, dict):
            if x.get("score100") is not None:
                return f"{x.get('score100')}/100"
            if x.get("total") is not None:
                detail = ", ".join(f"{k} {v}" for k, v in x.items() if k != "total" and not isinstance(v, (dict, list)))
                return f"{x.get('total')}" + (f" ({detail})" if detail else "")
            return "; ".join(f"{k} {v}" for k, v in list(x.items())[:5] if not isinstance(v, (dict, list)))
        return str(x) if x is not None else "N/A"
    if eod_flat.get("signalScore") is not None or eod_flat.get("v3FullScore") is not None:
        _v3("Signal score / V3 score", f"{_score_str(eod_flat.get('signalScore'))} / {_score_str(eod_flat.get('v3FullScore'))}", "Điểm hệ thống LHInvestment — đối chiếu với điểm tổng hợp mục 1B")
    if div:
        _v3("Phân kỳ RSI/MACD", div.get("type") or ("Phân kỳ dương" if div.get("bullish") else ("Phân kỳ âm" if div.get("bearish") else None)), str(div.get("message") or "")[:220])
    if ichi:
        _ichi_state = str(ichi.get("state") or "")
        _ichi_note = "giá dưới mây — xu hướng trung hạn chưa ủng hộ" if _ichi_state == "below_cloud" else ("giá trên mây — xu hướng trung hạn ủng hộ" if _ichi_state == "above_cloud" else "giá trong mây/chưa rõ xu hướng")
        _v3("Ichimoku", f"Tenkan {_fmt(ichi.get('tenkan'))}; Kijun {_fmt(ichi.get('kijun'))}; mây {_fmt(ichi.get('cloudBottom'))}–{_fmt(ichi.get('cloudTop'))}", f"Trạng thái {_ichi_state or 'N/A'}: {_ichi_note}")
    if _num(eod_flat.get("vwapDay")) is not None:
        _v3("VWAP ngày", _fmt(eod_flat.get("vwapDay")), "Giá so với VWAP phản ánh lực mua/bán bình quân phiên")
    _rr = eod_flat.get("riskReward")
    if isinstance(_rr, dict):
        _v3("Risk/Reward", "; ".join(f"{k}: {v}" for k, v in list(_rr.items())[:4]), "")
    elif _rr not in (None, ""):
        _v3("Risk/Reward", _rr, "")
    _fib = eod_flat.get("fibonacciLevelsDay")
    if isinstance(_fib, list) and _fib:
        _fib_nums = [_num(x) for x in _fib if _num(x) is not None]
        _nearest_note = "Các mốc cản/hồi Fibonacci gần nhất"
        if _num(close_v) is not None and _fib_nums:
            _near = min(_fib_nums, key=lambda x: abs(x - _num(close_v)))
            _nearest_note = f"Mốc gần giá nhất: {_fmt(_near)} ({_pct_delta(_near, close_v)} so với close {_fmt(close_v)}); các mốc còn lại là vùng tham chiếu."
        _v3("Fibonacci (ngày)", ", ".join(_fmt(x) for x in _fib[:6]), _nearest_note)
    if eod_flat.get("donchianHighDay") is not None or eod_flat.get("donchianLowDay") is not None:
        _v3("Donchian (ngày)", f"High {_fmt(eod_flat.get('donchianHighDay'))} / Mid {_fmt(eod_flat.get('donchianMidDay'))} / Low {_fmt(eod_flat.get('donchianLowDay'))}", "Kênh breakout 20 phiên")
    if v3_rows:
        _table(doc, ["Tín hiệu hệ thống", "Giá trị", "Diễn giải"], v3_rows)
    else:
        _bullets(doc, ["Cache V3 chưa có tín hiệu hệ thống cho mã này — cần refresh cache LHInvestment."])

    # 3D. Strategy - fixed LHinvt web strategy names/logic, not generic invented labels.
    _heading(doc, "3D. Strategy", 2)
    _near_support_pct = None
    _resistance_gap_pct = None
    try:
        if _num(close_v) is not None and _num(rs_s) is not None and _num(close_v) > 0:
            _near_support_pct = abs(_num(close_v) - _num(rs_s)) / _num(close_v) * 100
        if _num(close_v) is not None and _num(rs_r) is not None and _num(close_v) > 0:
            _resistance_gap_pct = (_num(rs_r) - _num(close_v)) / _num(close_v) * 100
    except Exception:
        pass
    _above_cloud = bool(ichi and str(ichi.get("state")) == "above_cloud")
    _below_cloud = bool(ichi and str(ichi.get("state")) == "below_cloud")
    _macd_cross_up = _num(macd_v) is not None and _num(sig_v) is not None and _num(macd_v) >= _num(sig_v)
    _macd_hist_ok = _num(hist_v) is not None and _num(hist_v) >= -0.02
    _macd_recover = _macd_hist_ok or _macd_cross_up
    _rsi_ok_pullback = _num(rsi_v) is not None and 48 <= _num(rsi_v) <= 62
    _rsi_low_watch = _num(rsi_v) is not None and _num(rsi_v) <= 45
    _near3 = _near_support_pct is not None and _near_support_pct <= 3.0
    _near25 = _near_support_pct is not None and _near_support_pct <= 2.5
    _vol_ok = _num(volratio_v) is not None and _num(volratio_v) >= 0.8
    _vol_breakout = _num(volratio_v) is not None and _num(volratio_v) >= 1.2
    _adx_ok = _num(adx_v) is not None and _num(adx_v) >= 20
    _di_bull = _num(pdi_v) is not None and _num(mdi_v) is not None and _num(pdi_v) >= _num(mdi_v)
    _ma20_ok = _num(close_v) is not None and _num(ma20) is not None and _num(close_v) >= _num(ma20)
    _ma50_ok = _num(close_v) is not None and _num(ma50) is not None and _num(close_v) >= _num(ma50)
    _ma200_ok = _num(close_v) is not None and _num(ma200) is not None and _num(close_v) >= _num(ma200)
    _bb_not_high = _num(bbp) is None or _num(bbp) <= 0.75
    _bb_rebound_zone = _num(bbp) is not None and _num(bbp) <= 0.55
    _bear_div = bool(div and (div.get("bearish") or str(div.get("type") or "").lower().find("bear") >= 0 or str(div.get("type") or "").lower().find("âm") >= 0))
    _bull_div = bool(div and (div.get("bullish") or str(div.get("type") or "").lower().find("bull") >= 0 or str(div.get("type") or "").lower().find("dương") >= 0))

    def _yn(x: bool) -> str:
        return "Đạt" if x else "Chưa đạt"

    evidence_common = (
        f"Close {_fmt(close_v)}; hỗ trợ {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)}), kháng cự {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)}); "
        f"MA20/50/200 {_fmt(ma20)}/{_fmt(ma50)}/{_fmt(ma200)}; Ichimoku {str(ichi.get('state') if ichi else 'N/A')}; "
        f"RSI14 {_fmt(rsi_v)}; MACD {_fmt(macd_v, 4)} / signal {_fmt(sig_v, 4)} / hist {_fmt(hist_v, 4)}; "
        f"Bollinger L/M/U {_fmt(bb_l)}/{_fmt(bb_m)}/{_fmt(bb_u)}, %B {_fmt(bbp)}; "
        f"VolumeRatio {_fmt(volratio_v)}, ADX {_fmt(adx_v)}, +DI/-DI {_fmt(pdi_v)}/{_fmt(mdi_v)}; "
        f"V3 zone {eod_flat.get('zoneState') or 'N/A'}, setup {eod_flat.get('setupType') or 'N/A'}, score {_score_str(eod_flat.get('signalScore') or eod_flat.get('v3FullScore'))}; "
        f"phân kỳ {'âm' if _bear_div else ('dương' if _bull_div else 'không rõ/không có')}.")
    _bullets(doc, ["3D dùng trực tiếp bộ chỉ báo chiến lược LHinvt: xu hướng MA/Ichimoku, hỗ trợ-kháng cự, RSI, MACD histogram/cross, Bollinger %B, volume/ADX/DI, phân kỳ và trạng thái V3. Không gắn nhãn BUY nếu thiếu xác nhận giá + dòng tiền."])

    strat_rows = []
    _trend_checks = [
        ("Giá trên mây Ichimoku", _above_cloud), ("gần hỗ trợ ≤3%", _near3), ("RSI 48-62", _rsi_ok_pullback),
        ("MACD/hist hồi", _macd_recover), ("volume không quá yếu", _vol_ok), ("không có phân kỳ âm", not _bear_div),
        ("không nằm vùng Bollinger quá cao", _bb_not_high),
    ]
    _trend_pass = sum(1 for _, ok in _trend_checks if ok)
    _trend_status = "WATCH/BUY khi có nến xác nhận" if _trend_pass >= 6 and _above_cloud and _near3 else ("Watch" if _trend_pass >= 4 else "Reject")
    strat_rows.append([
        "Trend Pullback Pro",
        _trend_status,
        "; ".join(f"{name}: {_yn(ok)}" for name, ok in _trend_checks),
        f"{evidence_common} Kết luận: {'ưu tiên chờ pullback giữ hỗ trợ rồi bật lên' if _trend_status != 'Reject' else 'chưa đủ nền pullback pro, không mua đuổi/không nâng tỷ trọng'}.",
    ])

    _support_checks = [
        ("gần hỗ trợ ≤2,5%", _near25), ("RSI thấp hoặc vùng hồi hợp lệ", _rsi_low_watch or (_above_cloud and _rsi_ok_pullback)),
        ("%B ≤0,55", _bb_rebound_zone), ("MACD histogram/cross đang hồi", _macd_recover),
        ("volume xác nhận", _vol_ok), ("không phân kỳ âm", not _bear_div),
    ]
    _support_pass = sum(1 for _, ok in _support_checks if ok)
    _support_status = "Watch yếu / chưa đủ BUY" if _support_pass >= 4 and _near25 else ("Reject/Watch" if _support_pass >= 3 else "Reject")
    strat_rows.append([
        "Support Rebound Hunter",
        _support_status,
        "; ".join(f"{name}: {_yn(ok)}" for name, ok in _support_checks),
        f"{evidence_common} Kết luận: chỉ chuyển BUY nếu có nến bật tại {_fmt(rs_s)} kèm VolumeRatio > 1,2 và MACD/RSI xác nhận; nếu mất hỗ trợ thì loại setup.",
    ])

    _shake_checks = [
        ("có breakdown dưới hỗ trợ 2-4,5%", False), ("reclaim hỗ trợ trong 1-3 nến", False),
        ("volume bán không tiếp diễn", _vol_ok), ("RSI/MACD hồi sau rũ", _macd_recover),
        ("không có phân kỳ âm", not _bear_div),
    ]
    strat_rows.append([
        "Shakeout Rebound",
        "Reject hiện tại",
        "; ".join(f"{name}: {_yn(ok)}" for name, ok in _shake_checks),
        f"Không gắn nhãn shakeout nếu chưa thấy breakdown-reclaim rõ quanh hỗ trợ {_fmt(rs_s)}. Cần kiểm tra thêm nến ngày/volume tại vùng rũ để tránh nhận nhầm breakdown thật.",
    ])

    _lh4_checks = [
        ("giá trên MA20", _ma20_ok), ("giá trên MA50", _ma50_ok), ("giá trên MA200", _ma200_ok),
        ("Ichimoku ủng hộ", _above_cloud and not _below_cloud), ("ADX ≥20", _adx_ok), ("+DI ≥ -DI", _di_bull),
        ("volume breakout ≥1,2", _vol_breakout), ("MACD/RSI đồng thuận", _macd_recover and (_num(rsi_v) is not None and _num(rsi_v) >= 50)),
    ]
    _lh4_pass = sum(1 for _, ok in _lh4_checks if ok)
    _lh4_status = "Watch" if _lh4_pass >= 6 and _vol_breakout else "Reject"
    strat_rows.append([
        "LH4 Wave Entry",
        _lh4_status,
        "; ".join(f"{name}: {_yn(ok)}" for name, ok in _lh4_checks),
        f"LH4 cần chất lượng sóng/xu hướng/dòng tiền đồng thuận. Hiện đạt {_lh4_pass}/{len(_lh4_checks)} điều kiện; {'chờ breakout qua kháng cự ' + _fmt(rs_r) if _lh4_status == 'Watch' else 'chưa đủ chất lượng để coi là điểm vào LH4'}.",
    ])
    _table(doc, ["Chiến lược LHinvt", "Kết luận", "Checklist PTKT", "Diễn giải hành động"], strat_rows)

    missing = []
    for must in ("ADX", "Ichimoku"):
        if must == "ADX" and _num(adx_v) is not None:
            continue
        if must == "Ichimoku" and ichi:
            continue
        if must.lower() not in analysis_text.lower() and all(must.lower() not in " ".join(str(x) for x in r).lower() for r in rows):
            missing.append(must)
    if missing:
        _bullets(doc, [f"Chỉ báo thiếu dữ liệu trong kỳ này: {', '.join(missing)} — cần refresh cache trước khi ra quyết định dựa trên nhóm chỉ báo này."])

    _heading(doc, "4. Vĩ mô / định giá / dữ liệu cơ bản", 1)
    macro_rows, macro_impact_rows = _macro_tables_for_symbol(symbol)
    _heading(doc, "4A. Vĩ mô hiện tại", 2)
    _table(doc, ["Vĩ mô hiện tại", "Giá trị mới nhất", "Mốc thời gian", "Nhận xét"], macro_rows)
    _heading(doc, "4B. Tác động đến cổ phiếu", 2)
    _table(doc, ["Yếu tố vĩ mô", "Diễn biến hiện tại/2026", "Cơ chế tác động", "Nhóm ngành/cổ phiếu liên quan"], macro_impact_rows)

    fundamental = _pick_lines(analysis_text, ("business", "doanh thu", "lợi nhuận", "EPS", "P/E", "P/B", "ROE", "ROA", "NIM", "target", "upside", "broker", "CTCK", "Yahoo", "24hmoney", "định giá"), 12)
    _, fsignal = load_json_record(symbol, [r"data\fundamental_signals.json", r"firebase_public\data\fundamental_signals.json"])
    _, fupside = load_json_record(symbol, [r"data\fundamental_top_upside.json", r"firebase_public\data\fundamental_top_upside.json", r"data\fa_market_valuation_breadth_summary.json", r"firebase_public\data\fa_market_valuation_breadth_summary.json"])
    fund_rows = []
    if isinstance(fsignal, dict):
        for k in ("pe", "pb", "eps", "roe", "roa", "revenueGrowth", "profitGrowth", "targetPrice", "upside", "rating", "sector", "marketCap"):
            if fsignal.get(k) not in (None, ""):
                fund_rows.append([k, fsignal.get(k)])
    # Map schema thật của fundamental_signals (cafef): broker/target_price/report_date/title, chuẩn hóa đơn vị VND → nghìn đồng.
    if isinstance(fsignal, dict):
        _close_n = _num(close_v)
        _tp_raw = _num(fsignal.get("target_price"))
        if _tp_raw:
            _tp_norm = _tp_raw / 1000 if (_close_n and _tp_raw / _close_n > 100) else _tp_raw
            fund_rows.append(["targetPrice", f"{_fmt(_tp_norm)} (CTCK {fsignal.get('broker') or 'N/A'}, ngày {fsignal.get('report_date') or 'N/A'})"])
            if _close_n:
                fund_rows.append(["upside", _pct_delta(_tp_norm, close_v)])
        for _k_src, _label in (("stop_loss", "Stop loss (CTCK)"), ("buy_low", "Vùng mua thấp (CTCK)"), ("buy_high", "Vùng mua cao (CTCK)")):
            _v2 = _num(fsignal.get(_k_src))
            if _v2:
                fund_rows.append([_label, _fmt(_v2 / 1000 if (_close_n and _v2 / _close_n > 100) else _v2)])
        if fsignal.get("title"):
            fund_rows.append(["Báo cáo CTCK mới nhất", str(fsignal.get("title"))[:180]])
    if isinstance(fupside, dict):
        for k in ("target_mean", "target_median", "targetMean", "targetMedian", "upside", "broker_count", "recommendation", "valuation", "score"):
            if fupside.get(k) not in (None, ""):
                fund_rows.append([k, fupside.get(k)])
    # Bổ sung chỉ số định giá từ Yahoo Finance (evidence-based, ghi rõ nguồn) khi cache thiếu P/E, P/B, ROE...
    import os as _os
    if _os.environ.get("SUPERLH_YAHOO_FUND", "1").lower() not in ("0", "false", "no"):
        try:
            from model3_lhinvestment_context import fetch_yahoo_fundamental
            _y = fetch_yahoo_fundamental(symbol) or {}
            _yd = _y.get("data") or {}
            def _yraw(module, key):
                v = (_yd.get(module) or {}).get(key)
                return v.get("raw") if isinstance(v, dict) else v
            for _lbl, _vv, _is_pct in (
                ("P/E (Yahoo)", _yraw("summaryDetail", "trailingPE"), False),
                ("P/B (Yahoo)", _yraw("defaultKeyStatistics", "priceToBook"), False),
                ("EPS (Yahoo)", _yraw("defaultKeyStatistics", "trailingEps"), False),
                ("ROE (Yahoo)", _yraw("financialData", "returnOnEquity"), True),
                ("Biên LN ròng (Yahoo)", _yraw("financialData", "profitMargins"), True),
                ("Tăng trưởng doanh thu (Yahoo)", _yraw("financialData", "revenueGrowth"), True),
                ("Vốn hóa (Yahoo)", _yraw("price", "marketCap"), False),
            ):
                _vn = _num(_vv)
                if _vn is None:
                    continue
                fund_rows.append([_lbl, f"{_vn * 100:.1f}%" if _is_pct else _fmt(_vn)])
        except Exception:
            pass
    targets = re.findall(r"(?:target(?: price)?|giá mục tiêu)[^\d]{0,30}([\d.,]+)", analysis_text, flags=re.I)
    numeric = []
    for t in targets:
        try:
            numeric.append(float(t.replace(".", "").replace(",", ".")))
        except Exception:
            pass
    # 4C: render dữ liệu fundamental/định giá thật thay vì chỉ đếm số dòng trong cache.
    _heading(doc, "4C. Chỉ tiêu cơ bản & định giá của cổ phiếu", 2)
    fund_label_map = {
        "pe": "P/E", "pb": "P/B", "eps": "EPS", "roe": "ROE", "roa": "ROA",
        "revenueGrowth": "Tăng trưởng doanh thu", "profitGrowth": "Tăng trưởng lợi nhuận",
        "targetPrice": "Giá mục tiêu", "upside": "Upside", "rating": "Xếp hạng",
        "sector": "Ngành", "marketCap": "Vốn hóa",
        "target_mean": "Target trung bình (CTCK)", "target_median": "Target trung vị (CTCK)",
        "targetMean": "Target trung bình (CTCK)", "targetMedian": "Target trung vị (CTCK)",
        "broker_count": "Số CTCK theo dõi", "recommendation": "Khuyến nghị", "valuation": "Định giá", "score": "Điểm",
    }
    fund_table_rows = [[fund_label_map.get(str(k), str(k)), str(v) if (isinstance(v, str) and ("%" in v or not v.replace(".", "").replace(",", "").replace("+", "").replace("-", "").isdigit())) else _fmt(v)] for k, v in fund_rows]
    backend_reports = _backend_report_lines(symbol)
    ratio_sector_rows = _ratio_sector_rows(symbol)
    if numeric:
        numeric.sort()
        mean = sum(numeric) / len(numeric)
        median = numeric[len(numeric)//2] if len(numeric) % 2 else (numeric[len(numeric)//2-1] + numeric[len(numeric)//2]) / 2
        fund_table_rows.append(["Target từ phân tích (n=%d)" % len(numeric), f"trung bình {_fmt(mean)}; trung vị {_fmt(median)}; upside so với giá tham chiếu {_pct_delta(mean, close_v)}"])
    if not fund_table_rows:
        fund_table_rows.append(["Dữ liệu cơ bản", "Chưa có chỉ tiêu định lượng đã kiểm chứng trong kỳ này — phần định giá cần bổ sung thủ công trước khi ra quyết định; không suy đoán số liệu."])
    _table(doc, ["Chỉ tiêu", "Giá trị"], fund_table_rows)
    if ratio_sector_rows:
        _heading(doc, "4D. So sánh P/E, P/B, P/S với nhóm ngành", 2)
        _table(doc, ["Chỉ tiêu", "Mã cổ phiếu", "Nhóm ngành", "Trung bình ngành", "Nhận xét"], ratio_sector_rows)
    if fundamental:
        _bullets(doc, _dedupe_lines(fundamental, seen_keys)[:6])

    _heading(doc, "5. Kịch bản đầu tư & hành động chính", 1)
    scenario_rows = [
        ["Bull", f"Kích hoạt khi giá vượt kháng cự {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)} so với giá tham chiếu {_fmt(close_v)}) với VolumeRatio > 1.2, RSI giữ trên 50 và MACD/histogram xác nhận. Hành động: chỉ nâng tỷ trọng sau breakout rõ, tránh mua đuổi trước tín hiệu."],
        ["Base", f"Giá dao động trong vùng hỗ trợ {_fmt(rs_s)} - kháng cự {_fmt(rs_r)}; tin tích cực đã có nhưng cần xác nhận bằng dòng tiền. Hành động: theo dõi phản ứng tại hỗ trợ/kháng cự, ưu tiên chờ nến xác nhận."],
        ["Bear", f"Giá mất hỗ trợ {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)} so với giá), thanh khoản bán tăng, RSI/MACD yếu đi hoặc xuất hiện tin bất lợi về lợi nhuận/pha loãng. Hành động: giảm rủi ro, không bắt đáy khi chưa có đảo chiều."],
        ["Catalyst", "Theo dõi KQKD quý, kế hoạch lợi nhuận, tăng vốn/ESOP/cổ tức, cập nhật target CTCK, thị phần môi giới, dư nợ margin và phản ứng giá tại vùng hỗ trợ/kháng cự."],
    ]
    extracted_scenarios = _pick_lines(scenario_text or analysis_text, ("Bull", "Base", "Bear", "Catalyst", "kịch bản", "trigger", "kích hoạt", "invalidation", "xác suất", "vùng giá", "giải ngân", "theo dõi"), 10)
    for line in _dedupe_lines(extracted_scenarios, seen_keys)[:4]:
        scenario_rows.append(["Bổ sung", line])
    _table(doc, ["Kịch bản", "Điều kiện - tác động - hành động"], scenario_rows)

    _heading(doc, "6. Rủi ro & quan điểm", 1)
    risk_items = _pick_lines(risk_text, ("Risk", "rủi ro", "Invalidation", "Stance", "WATCH", "AVOID", "stop", "thanh khoản", "khuyến nghị", "theo dõi"), 14)
    # Risk score định lượng 1-5: mỗi tiêu chí rủi ro cộng 1 điểm, nêu rõ căn cứ.
    risk_reasons = []
    if _num(volratio_v) is None or _num(volratio_v) < 0.8:
        risk_reasons.append(f"VolumeRatio {_fmt(volratio_v)} thiếu/thấp, tín hiệu dòng tiền chưa đáng tin.")
    if _num(rsi_v) is not None and _num(rsi_v) < 50:
        risk_reasons.append(f"RSI {_fmt(rsi_v)} dưới 50, động lượng chưa xác nhận xu hướng tăng.")
    if _num(adx_v) is not None and _num(adx_v) < 20:
        risk_reasons.append(f"ADX {_fmt(adx_v)} dưới 20, xu hướng yếu/dễ nhiễu.")
    elif _num(adx_v) is None:
        risk_reasons.append("ADX thiếu dữ liệu, chưa đo được sức mạnh xu hướng.")
    if _num(close_v) is not None and _num(ma200) is not None and _num(close_v) < _num(ma200):
        risk_reasons.append(f"Giá dưới MA200 ({_fmt(ma200)}), xu hướng dài hạn chưa ủng hộ.")
    # Các cảnh báo hệ thống V3 phải được tính vào rủi ro, không chỉ hiển thị ở 3C.
    if isinstance(eod_flat.get("zoneState"), str) and "quá mua" in str(eod_flat.get("zoneState")).lower():
        risk_reasons.append("Zone hệ thống V3: Quá mua — hạn chế mua đuổi.")
    if ichi and str(ichi.get("state")) == "below_cloud":
        risk_reasons.append("Giá dưới mây Ichimoku — xu hướng trung hạn chưa ủng hộ.")
    if _num(pdi_v) is not None and _num(mdi_v) is not None and _num(mdi_v) > _num(pdi_v):
        risk_reasons.append(f"-DI {_fmt(mdi_v)} trên +DI {_fmt(pdi_v)} — bên bán đang chiếm ưu thế.")
    if len(news_lines) == 0:
        risk_reasons.append("Không có tin đã kiểm chứng trong kỳ — rủi ro thiếu thông tin sự kiện.")
    if str(close_v) == "N/A" or str(rs_s) == "N/A" or str(rs_r) == "N/A":
        risk_reasons.append("Thiếu close/hỗ trợ/kháng cự để định lượng vùng hành động.")
    risk_level = min(5, max(1, 1 + len(risk_reasons)))
    risk_score = f"{risk_level}/5"
    _table(doc, ["Hạng mục rủi ro", "Dữ liệu cụ thể", "Quan điểm"], [
        ["Risk score", f"{risk_score} ({len(risk_reasons)} yếu tố rủi ro)", "; ".join(risk_reasons) or "Rủi ro thấp theo tiêu chí định lượng; vẫn cần xác nhận bằng giá và volume."],
        ["Vùng giá", f"Close {_fmt(close_v)}; hỗ trợ {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)} so với giá); kháng cự {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)} so với giá); stop/invalid {(rs.get('stopLossDay') if isinstance(rs, dict) else None) or 'N/A'}", "Không mua đuổi sát kháng cự nếu volume không xác nhận; mất hỗ trợ thì giảm rủi ro."],
        ["Động lượng", f"RSI14 {_fmt(rsi_v)}; MACD {_fmt(macd_v, 4)}; signal {_fmt(sig_v, 4)}; hist {_fmt(hist_v, 4)}; ROC20 {_fmt(val('roc20','ROC20'))}", "Ưu tiên khi RSI > 50 và MACD mở rộng; tránh khi histogram âm/mỏng và ROC yếu."],
        ["Dòng tiền", f"Volume {_fmt(vol_v)}; AvgVol20 {_fmt(avgvol_v)}; VolumeRatio {_fmt(volratio_v)}; ADX {_fmt(adx_v)}; +DI {_fmt(pdi_v)}; -DI {_fmt(mdi_v)}", "Breakout không có volume hoặc ADX yếu dễ là bull trap."],
        ["Rủi ro dữ liệu", f"Chỉ tiêu cơ bản: {len(fund_rows)}; target CTCK phát hiện: {len(numeric)}; tin đã kiểm chứng: {len(news_lines)}", "Nếu target/news/fundamental thiếu, báo cáo chỉ dùng để watchlist, cần manual review."],
    ])
    _bullets(doc, _dedupe_lines(risk_items, seen_keys) or ["Nội dung rủi ro bổ sung chưa đủ chuẩn kiểm chứng trong kỳ này; bảng rủi ro trên dùng dữ liệu LHInvestment, indicator và số tin đã xác thực để đánh giá."])

    _heading(doc, "7. Kế hoạch theo dõi", 1)
    followup_items = _pick_lines(followup_text, ("trade plan", "kế hoạch", "điều kiện", "hỗ trợ", "kháng cự", "cắt", "mua", "volume", "MACD", "RSI", "refresh", "trigger"), 14)
    _table(doc, ["Việc cần theo dõi", "Ngưỡng/số liệu cụ thể", "Hành động"], [
        ["Giá & vùng kỹ thuật", f"Close {_fmt(close_v)}; hỗ trợ {_fmt(rs_s)} ({_pct_delta(rs_s, close_v)}); kháng cự {_fmt(rs_r)} ({_pct_delta(rs_r, close_v)}); Bollinger {_fmt(bb_l)}/{_fmt(bb_m)}/{_fmt(bb_u)}; %B {_fmt(bbp)}", "Cập nhật mỗi phiên; chỉ quan sát mua khi giữ hỗ trợ hoặc breakout kháng cự có xác nhận."],
        ["Động lượng", f"RSI14 {_fmt(rsi_v)}; MACD {_fmt(macd_v, 4)}; signal {_fmt(sig_v, 4)}; hist {_fmt(hist_v, 4)}", "Trigger tích cực: RSI > 50 và MACD cắt lên/mở rộng; trigger tiêu cực: RSI mất 45 hoặc MACD xấu đi."],
        ["Dòng tiền", f"Volume {_fmt(vol_v)}; AvgVol20 {_fmt(avgvol_v)}; VolumeRatio {_fmt(volratio_v)}; ADX {_fmt(adx_v)}; +DI {_fmt(pdi_v)}; -DI {_fmt(mdi_v)}", "Yêu cầu VolumeRatio > 1.2 khi breakout; nếu volume thấp thì không nâng tỷ trọng."],
        ["Tin tức/catalyst", f"Tin đã kiểm chứng trong báo cáo: {len(news_lines)}; target CTCK phát hiện: {len(numeric)}", "Theo dõi KQKD, target CTCK, cổ tức/ESOP/phát hành thêm, IPO/M&A/sản phẩm/chuỗi mới."],
        ["Checklist hành động", "Bull/Base/Bear phải khớp cả news + kỹ thuật + volume", "Ra quyết định theo trigger đã xác nhận; nếu thiếu dữ liệu thì ghi manual review thay vì suy đoán."],
    ])
    _bullets(doc, _dedupe_lines(followup_items, seen_keys) or ["Nội dung kế hoạch bổ sung chưa đủ chuẩn kiểm chứng trong kỳ này; bảng kế hoạch trên dùng dữ liệu LHInvestment và trigger định lượng để không bỏ trống mục 8."])
    _bullets(doc, ["Báo cáo nghiên cứu dữ liệu; nhà đầu tư tự chịu trách nhiệm với quyết định giao dịch."])

    for para in doc.paragraphs:
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.05
        for run in para.runs:
            _force_run_font(run, font_name=VI_FONT)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _force_run_font(run, font_name=VI_FONT)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return str(path)
